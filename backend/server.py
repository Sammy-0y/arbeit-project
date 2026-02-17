from fastapi import FastAPI, APIRouter, HTTPException, Depends, status, UploadFile, File, Form, BackgroundTasks
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
import io
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr, ConfigDict, field_validator
from typing import Optional, Literal, List
from datetime import datetime, timezone, timedelta
import bcrypt
import jwt
import uuid
import re
import json
import secrets
#from emergentintegrations.llm.chat import LlmChat, UserMessage
from openai import AsyncOpenAI

# PDF and DOCX parsing
import pdfplumber
from docx import Document as DocxDocument

# Import notification service
from backend.notification_service import (
    send_email,
    get_new_job_email_template,
    get_candidate_status_change_email_template,
    get_interview_booked_email_template,
    send_client_user_welcome_email
)

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'arbeit-secret-key-change-in-production')
JWT_ALGORITHM = 'HS256'
JWT_EXPIRATION_HOURS = 24

# Security
security = HTTPBearer()


# ============ NOTIFICATION HELPER FUNCTIONS ============

async def send_candidate_status_change_notification(
    candidate_id: str,
    old_status: str,
    new_status: str,
    changed_by: str
):
    """Send notification when candidate status changes"""
    try:
        # Get candidate details
        candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
        if not candidate:
            return
        
        # Get job and client details
        job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
        if not job:
            return
            
        client_doc = await db.clients.find_one({"client_id": job["client_id"]}, {"_id": 0})
        if not client_doc:
            return
        
        # Generate email content
        subject, body = get_candidate_status_change_email_template(
            candidate=candidate,
            job=job,
            client=client_doc,
            new_status=new_status,
            changed_by=changed_by
        )
        
        # Get recruiters to notify
        recruiters = await db.users.find(
            {"role": {"$in": ["admin", "recruiter"]}},
            {"_id": 0, "email": 1}
        ).to_list(100)
        
        # Send to each recruiter
        for recruiter in recruiters:
            await send_email(recruiter["email"], subject, body)
        
        # Create in-app notification
        notification_doc = {
            "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
            "type": "candidate_status_change",
            "title": f"Candidate Status Changed: {candidate.get('name', 'Unknown')}",
            "message": f"Status changed from {old_status} to {new_status} by {changed_by}",
            "entity_type": "candidate",
            "entity_id": candidate_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "read": False,
            "recipients": ["admin", "recruiter"]
        }
        await db.notifications.insert_one(notification_doc)
        
    except Exception as e:
        print(f"Failed to send status change notification: {e}")


async def send_interview_booking_notification(
    interview_id: str,
    candidate_id: str,
    slot_time: str,
    booked_by: str
):
    """Send notification when interview slot is booked"""
    try:
        # Get interview details
        interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
        if not interview:
            return
        
        # Get candidate details
        candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
        if not candidate:
            return
        
        # Get job and client details
        job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
        if not job:
            return
            
        client_doc = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
        if not client_doc:
            return
        
        # Generate email content
        subject, body = get_interview_booked_email_template(
            interview=interview,
            candidate=candidate,
            job=job,
            client=client_doc,
            slot_time=slot_time
        )
        
        # Get recruiters to notify
        recruiters = await db.users.find(
            {"role": {"$in": ["admin", "recruiter"]}},
            {"_id": 0, "email": 1}
        ).to_list(100)
        
        # Send to each recruiter
        for recruiter in recruiters:
            await send_email(recruiter["email"], subject, body)
        
        # Get client users to notify
        client_users = await db.users.find(
            {"role": "client_user", "client_id": interview["client_id"]},
            {"_id": 0, "email": 1}
        ).to_list(100)
        
        for client_user in client_users:
            await send_email(client_user["email"], subject, body)
        
        # Create in-app notification
        notification_doc = {
            "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
            "type": "interview_booked",
            "title": f"Interview Booked: {candidate.get('name', 'Unknown')}",
            "message": f"Interview scheduled for {slot_time}",
            "entity_type": "interview",
            "entity_id": interview_id,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "read": False,
            "recipients": ["admin", "recruiter", interview["client_id"]]
        }
        await db.notifications.insert_one(notification_doc)
        
    except Exception as e:
        print(f"Failed to send interview booking notification: {e}")

# Create the main app
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# ============ MODELS ============

class UserRole(str):
    ADMIN = "admin"
    RECRUITER = "recruiter"
    CLIENT_USER = "client_user"

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    role: Literal["admin", "recruiter", "client_user"]
    client_id: Optional[str] = None

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class UserResponse(BaseModel):
    email: str
    name: str
    role: str
    client_id: Optional[str] = None
    phone: Optional[str] = None
    user_id: Optional[str] = None

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse
    must_change_password: Optional[bool] = False

class UserInDB(BaseModel):
    email: str
    name: str
    role: str
    client_id: Optional[str] = None
    password_hash: str
    created_at: str


# ============ CANDIDATE PORTAL MODELS ============

class CandidatePortalRegister(BaseModel):
    """Candidate registration for the portal"""
    email: EmailStr
    password: str
    name: str
    phone: str
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    experience_years: Optional[int] = None

class CandidatePortalLogin(BaseModel):
    """Candidate login"""
    email: EmailStr
    password: str

class CandidatePortalResponse(BaseModel):
    """Candidate portal user response"""
    candidate_portal_id: str
    email: str
    name: str
    phone: str
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    experience_years: Optional[int] = None
    created_at: str
    must_change_password: bool = False

class CandidatePortalTokenResponse(BaseModel):
    """Token response for candidate login"""
    access_token: str
    token_type: str
    candidate: CandidatePortalResponse
    must_change_password: bool = False

class CandidatePasswordChange(BaseModel):
    """Password change request for candidate portal"""
    current_password: str
    new_password: str

class SendSelectionNotificationRequest(BaseModel):
    """Request to send selection notification to candidate"""
    candidate_id: str
    custom_message: Optional[str] = None


# Phase 2: Client Management Models
class ClientCreate(BaseModel):
    company_name: str
    status: Literal["active", "inactive"] = "active"
    industry: Optional[str] = None
    website: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    country: Optional[str] = None
    postal_code: Optional[str] = None
    notes: Optional[str] = None

class ClientUpdate(BaseModel):
    company_name: Optional[str] = None
    status: Optional[Literal["active", "inactive"]] = None
    industry: Optional[str] = None
    website: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    country: Optional[str] = None
    postal_code: Optional[str] = None
    notes: Optional[str] = None

class ClientResponse(BaseModel):
    client_id: str
    company_name: str
    status: str
    created_at: str
    user_count: Optional[int] = 0
    industry: Optional[str] = None
    website: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    city: Optional[str] = None
    state: Optional[str] = None
    country: Optional[str] = None
    postal_code: Optional[str] = None
    notes: Optional[str] = None

class ClientUserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str
    phone: Optional[str] = None

class ClientUserUpdate(BaseModel):
    email: Optional[EmailStr] = None
    name: Optional[str] = None
    phone: Optional[str] = None

# Phase 3: Job Requirements Models
class ExperienceRange(BaseModel):
    min_years: int = Field(ge=0)
    max_years: int = Field(ge=0)
    
    @field_validator('max_years')
    @classmethod
    def validate_max(cls, max_years, info):
        min_years = info.data.get('min_years', 0)
        if max_years < min_years:
            raise ValueError('max_years must be >= min_years')
        return max_years

class SalaryRange(BaseModel):
    min_amount: Optional[int] = None
    max_amount: Optional[int] = None
    currency: str = "USD"
    
    @field_validator('max_amount')
    @classmethod
    def validate_max(cls, max_amount, info):
        min_amount = info.data.get('min_amount')
        if max_amount and min_amount and max_amount < min_amount:
            raise ValueError('max_amount must be >= min_amount')
        return max_amount

class JobCreate(BaseModel):
    title: str
    location: str
    employment_type: Literal["Full-time", "Part-time", "Contract"]
    experience_range: ExperienceRange
    salary_range: Optional[SalaryRange] = None
    work_model: Literal["Onsite", "Hybrid", "Remote"]
    required_skills: list[str] = []
    description: str
    status: Literal["Draft", "Active", "Closed"] = "Active"
    client_id: Optional[str] = None
    city: Optional[str] = None  # Mandatory for Onsite/Hybrid
    notice_period_days: Optional[int] = Field(None, ge=0)  # Notice period in days
    
    @field_validator('city')
    @classmethod
    def validate_city(cls, city, info):
        work_model = info.data.get('work_model')
        if work_model in ["Onsite", "Hybrid"] and not city:
            raise ValueError('City is mandatory for Onsite and Hybrid work models')
        return city
    
    @field_validator('notice_period_days')
    @classmethod
    def validate_notice_period(cls, notice_period_days):
        if notice_period_days is not None:
            allowed = [0, 7, 15, 30, 45, 60, 90]
            if notice_period_days not in allowed:
                raise ValueError(f'notice_period_days must be one of {allowed}')
        return notice_period_days

class JobUpdate(BaseModel):
    title: Optional[str] = None
    location: Optional[str] = None
    employment_type: Optional[Literal["Full-time", "Part-time", "Contract"]] = None
    experience_range: Optional[ExperienceRange] = None
    salary_range: Optional[SalaryRange] = None
    work_model: Optional[Literal["Onsite", "Hybrid", "Remote"]] = None
    required_skills: Optional[list[str]] = None
    description: Optional[str] = None
    status: Optional[Literal["Draft", "Active", "Closed"]] = None
    city: Optional[str] = None
    notice_period_days: Optional[int] = None

class JobResponse(BaseModel):
    job_id: str
    client_id: str
    title: str
    location: str
    employment_type: str
    experience_range: ExperienceRange
    salary_range: Optional[SalaryRange] = None
    work_model: str
    required_skills: list[str]
    description: str
    status: str
    created_at: str
    created_by: str
    company_name: Optional[str] = None  # Populated from client lookup

# Phase 4: Candidate Models
class ParsedResume(BaseModel):
    name: str = "Candidate"
    current_role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    skills: list[str] = []
    experience: list[dict] = []
    education: list[dict] = []
    summary: Optional[str] = None

class CandidateStory(BaseModel):
    headline: str = "Candidate Profile"
    summary: str = "Profile summary"
    timeline: list[dict] = []
    skills: list[str] = []
    fit_score: int = 0
    highlights: list[str] = []

class CandidateCreate(BaseModel):
    job_id: str
    name: str
    current_role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    skills: list[str] = []
    experience: list[dict] = []
    education: list[dict] = []
    summary: Optional[str] = None

class CandidateUpdate(BaseModel):
    name: Optional[str] = None
    current_role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    skills: Optional[list[str]] = None
    experience: Optional[list[dict]] = None
    education: Optional[list[dict]] = None
    summary: Optional[str] = None
    status: Optional[Literal["NEW", "PIPELINE", "APPROVED", "REJECTED"]] = None

class CandidateResponse(BaseModel):
    candidate_id: str
    job_id: str
    name: str
    current_role: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    linkedin: Optional[str] = None
    skills: list[str]
    experience: list[dict]
    education: list[dict]
    summary: Optional[str] = None
    cv_file_url: Optional[str] = None
    ai_story: Optional[CandidateStory] = None
    status: str
    created_at: str
    created_by: str

# Phase 5: Review Workflow Models
class ReviewAction(str):
    APPROVE = "APPROVE"
    PIPELINE = "PIPELINE"
    REJECT = "REJECT"
    COMMENT = "COMMENT"

class ReviewCreate(BaseModel):
    action: Literal["APPROVE", "PIPELINE", "REJECT", "COMMENT"]
    comment: Optional[str] = None

class ReviewResponse(BaseModel):
    review_id: str
    candidate_id: str
    user_id: str
    user_name: str
    user_role: str
    timestamp: str
    action: str
    comment: Optional[str] = None


# ============ GOVERNANCE MODELS (RBAC + Audit) ============

# RBAC: Client Roles with Permissions
class PermissionSet(BaseModel):
    # Job permissions
    can_view_jobs: bool = True
    can_create_jobs: bool = False
    can_edit_jobs: bool = False
    can_delete_jobs: bool = False
    
    # Candidate permissions
    can_view_candidates: bool = True
    can_create_candidates: bool = False
    can_edit_candidates: bool = False
    can_delete_candidates: bool = False
    
    # Candidate actions
    can_update_candidate_status: bool = False
    can_upload_cv: bool = False
    can_replace_cv: bool = False
    can_regenerate_story: bool = False
    
    # CV viewing permissions
    can_view_full_cv: bool = False
    can_view_redacted_cv: bool = True
    
    # Governance permissions
    can_view_audit_log: bool = False
    can_manage_roles: bool = False
    can_manage_users: bool = False
    can_export_reports: bool = False

class ClientRoleCreate(BaseModel):
    name: str
    description: Optional[str] = None
    permissions: PermissionSet

class ClientRoleUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    permissions: Optional[PermissionSet] = None

class ClientRoleResponse(BaseModel):
    role_id: str
    client_id: str
    name: str
    description: Optional[str] = None
    permissions: PermissionSet
    created_at: str
    updated_at: str

# User-Role Assignment
class UserRoleAssignment(BaseModel):
    user_id: str
    client_role_id: str

class UserRoleResponse(BaseModel):
    assignment_id: str
    user_id: str
    user_email: str
    client_id: str
    client_role_id: str
    role_name: str
    assigned_by: str
    created_at: str

# Audit Log
class AuditLogEntry(BaseModel):
    log_id: str
    timestamp: str
    user_id: str
    user_email: str
    user_role: str
    client_id: Optional[str] = None
    action_type: str
    entity_type: str
    entity_id: Optional[str] = None
    previous_value: Optional[dict] = None
    new_value: Optional[dict] = None
    metadata: Optional[dict] = None
    ip_address: Optional[str] = None

class AuditLogFilter(BaseModel):
    client_id: Optional[str] = None
    user_id: Optional[str] = None
    action_type: Optional[str] = None
    entity_type: Optional[str] = None
    entity_id: Optional[str] = None
    from_date: Optional[str] = None
    to_date: Optional[str] = None
    limit: int = 100
    skip: int = 0

# Access Matrix Response
class UserPermissionMatrix(BaseModel):
    user_id: str
    user_email: str
    user_name: str
    client_id: str
    client_name: str
    roles: list[str]
    permissions: PermissionSet


# ============ CV VERSIONING MODELS ============

class CVVersionCreate(BaseModel):
    candidate_id: str
    file_url: str
    source_filename: str
    ai_parsed_data: Optional[dict] = None
    ai_story_json: Optional[dict] = None
    fit_score: Optional[float] = None

class CVVersionResponse(BaseModel):
    version_id: str
    candidate_id: str
    version_number: int
    file_url: str
    source_filename: str
    uploaded_by_user_id: str
    uploaded_by_email: str
    uploaded_at: str
    is_active: bool
    ai_parsed_data: Optional[dict] = None
    ai_story_json: Optional[dict] = None
    fit_score: Optional[float] = None
    deleted_at: Optional[str] = None
    delete_type: Optional[str] = None  # None | "soft" | "hard"
    deleted_by_user_id: Optional[str] = None

class CVVersionListItem(BaseModel):
    version_id: str
    version_number: int
    source_filename: str
    uploaded_by_email: str
    uploaded_at: str
    is_active: bool
    deleted_at: Optional[str] = None
    delete_type: Optional[str] = None


# ============ INTERVIEW ORCHESTRATION MODELS ============

NOTICE_PERIOD_OPTIONS = [0, 7, 15, 30, 45, 60, 90]  # Days

class InterviewSlot(BaseModel):
    """Individual time slot for interview"""
    slot_id: str
    start_time: datetime
    end_time: datetime
    duration_minutes: int
    is_available: bool = True

class InterviewCreate(BaseModel):
    """Create new interview - client proposes slots"""
    job_id: str
    candidate_id: str
    interview_mode: Literal["Video", "Phone", "Onsite"]
    interview_duration: int = Field(ge=15, le=240, description="Duration in minutes")
    time_zone: str = Field(default="Asia/Kolkata")
    proposed_slots: List[dict]  # List of {start_time, end_time}
    meeting_link: Optional[str] = None
    additional_instructions: Optional[str] = None
    interview_round: int = Field(default=1, ge=1, le=10, description="Interview round number")
    round_name: Optional[str] = None  # e.g., "Technical Round", "HR Round"

class InterviewUpdate(BaseModel):
    """Update interview details"""
    interview_mode: Optional[Literal["Video", "Phone", "Onsite"]] = None
    meeting_link: Optional[str] = None
    additional_instructions: Optional[str] = None
    interview_status: Optional[Literal[
        "Draft",
        "Awaiting Candidate Confirmation",
        "Confirmed",
        "Scheduled",
        "Completed",
        "No Show",
        "Cancelled",
        "Passed",
        "Failed"
    ]] = None
    feedback: Optional[str] = None
    rating: Optional[int] = Field(default=None, ge=1, le=5)

class CandidateSlotSelection(BaseModel):
    """Candidate selects a slot"""
    slot_id: str
    confirmed: bool = True

class InterviewResponse(BaseModel):
    """Interview response model"""
    interview_id: str
    job_id: str
    candidate_id: str
    client_id: str
    candidate_name: Optional[str] = None
    job_title: Optional[str] = None
    company_name: Optional[str] = None
    interview_mode: str
    interview_duration: int
    scheduled_start_time: Optional[str] = None
    scheduled_end_time: Optional[str] = None
    time_zone: str
    interview_status: str
    meeting_link: Optional[str] = None
    additional_instructions: Optional[str] = None
    invite_sent: bool
    invite_sent_by: Optional[str] = None
    candidate_confirmation_timestamp: Optional[str] = None
    no_show_flag: bool
    no_show_count: int = 0
    proposed_slots: List[dict]
    selected_slot_id: Optional[str] = None
    created_at: str
    updated_at: str
    created_by: str
    interview_round: int = 1
    round_name: Optional[str] = None
    feedback: Optional[str] = None
    rating: Optional[int] = None

class InterviewListItem(BaseModel):
    """Lightweight interview item for lists"""
    interview_id: str
    job_id: str
    candidate_id: str
    candidate_name: Optional[str] = None
    job_title: Optional[str] = None
    interview_mode: str
    interview_status: str
    scheduled_start_time: Optional[str] = None
    created_at: str
    interview_round: int = 1
    round_name: Optional[str] = None

class InterviewPipelineStats(BaseModel):
    """Interview pipeline statistics for dashboard"""
    total_interviews: int
    awaiting_confirmation: int
    confirmed: int
    scheduled: int
    completed: int
    no_shows: int
    cancelled: int


# ============ UTILITIES ============

def hash_password(password: str) -> str:
    """Hash a password using bcrypt"""
    salt = bcrypt.gensalt()
    hashed = bcrypt.hashpw(password.encode('utf-8'), salt)
    return hashed.decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password against its hash"""
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

# Phase 4: File storage setup
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = BASE_DIR / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


async def save_cv_file(file: UploadFile, candidate_id: str) -> str:
    """Save uploaded CV file and return URL"""
    file_extension = Path(file.filename).suffix
    filename = f"{candidate_id}{file_extension}"
    file_path = UPLOAD_DIR / filename
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    return f"/api/uploads/{filename}"


async def extract_text_from_cv(file: UploadFile) -> str:
    """Extract text from CV file (PDF, DOCX, or plain text)"""
    await file.seek(0)
    file_content = await file.read()
    filename = file.filename.lower()
    
    extracted_text = ""
    
    try:
        if filename.endswith('.pdf'):
            # Extract text from PDF using pdfplumber
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += page_text + "\n"
            print(f"[DEBUG] Extracted {len(extracted_text)} chars from PDF")
            
        elif filename.endswith('.docx'):
            # Extract text from DOCX
            doc = DocxDocument(io.BytesIO(file_content))
            for paragraph in doc.paragraphs:
                extracted_text += paragraph.text + "\n"
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        extracted_text += cell.text + " "
                    extracted_text += "\n"
            print(f"[DEBUG] Extracted {len(extracted_text)} chars from DOCX")
            
        elif filename.endswith('.doc'):
            # For .doc files, try basic decoding
            try:
                extracted_text = file_content.decode('utf-8', errors='ignore')
            except:
                extracted_text = file_content.decode('latin-1', errors='ignore')
            print(f"[DEBUG] Extracted {len(extracted_text)} chars from DOC (basic)")
            
        elif filename.endswith(('.txt', '.rtf')):
            # Plain text files
            extracted_text = file_content.decode('utf-8', errors='ignore')
            print(f"[DEBUG] Extracted {len(extracted_text)} chars from text file")
            
        else:
            # Try UTF-8 decoding as fallback
            extracted_text = file_content.decode('utf-8', errors='ignore')
            print(f"[DEBUG] Fallback text extraction: {len(extracted_text)} chars")
    
    except Exception as e:
        print(f"[ERROR] Text extraction failed: {e}")
        extracted_text = f"CV Upload - {file.filename}"
    
    # Clean up extracted text
    if extracted_text:
        # Remove excessive whitespace but preserve structure
        extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
        extracted_text = re.sub(r' {2,}', ' ', extracted_text)
        extracted_text = extracted_text.strip()
    
    return extracted_text if extracted_text else f"CV Upload - {file.filename}"

def redact_text(text: str) -> str:
    """Redact personal information from text"""
    # Email redaction
    text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL REDACTED]', text)
    
    # Phone redaction (various formats including short forms)
    text = re.sub(r'\b(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}\b', '[PHONE REDACTED]', text)
    text = re.sub(r'\b\d{3}-\d{4}\b', '[PHONE REDACTED]', text)  # Short format like 555-1234
    
    # LinkedIn URLs
    text = re.sub(r'https?://(www\.)?linkedin\.com/[^\s]+', '[LINKEDIN REDACTED]', text)
    
    # Generic URLs (potential personal sites)
    text = re.sub(r'https?://[^\s]+', '[URL REDACTED]', text)
    
    return text

async def call_openai_directly(system_prompt: str, user_prompt: str, api_key: str) -> str:
    """Call OpenAI API directly using the official SDK"""
    try:
        client = AsyncOpenAI(api_key=api_key)
        response = await client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        raise Exception(f"OpenAI API call failed: {str(e)}")

async def parse_cv_with_ai(cv_text: str, existing_data: dict = None) -> ParsedResume:
    """Parse CV using RecruitAssist AI with enhanced extraction"""
    llm_key = os.environ.get('EMERGENT_LLM_KEY')
    if not llm_key:
        # Return fallback data if no LLM key
        if existing_data:
            return ParsedResume(**existing_data)
        return ParsedResume(
            name="CV Upload",
            summary="AI parsing unavailable - please edit manually"
        )
    
    try:
        # Enhanced RecruitAssist AI System Prompt
        system_prompt = """You are an expert CV/Resume parser. Extract ALL information from the resume text.

CRITICAL CONTACT EXTRACTION - DO NOT MISS:
1. EMAIL: Look for @ symbol anywhere in the text (e.g., name@gmail.com, user@company.co.in)
2. PHONE: Look for 10+ digit numbers, possibly with +91, country codes, spaces, dashes
3. LINKEDIN: Look for linkedin.com/in/ URLs

SEARCH THESE COMMON LOCATIONS FOR CONTACT INFO:
- Header/top section
- Footer/bottom section  
- "Contact" or "Personal Details" sections
- Near the candidate name

EXPERIENCE EXTRACTION RULES:
- DEDUPLICATE: If the same company appears multiple times, MERGE into a single entry with full date range
- Each company should appear ONLY ONCE in the experience list
- If candidate had multiple roles at same company, combine them with the full tenure duration

Return ONLY valid JSON with this exact structure:

{
  "name": "FULL NAME - extract the candidate's complete name",
  "current_role": "Most recent job title/designation",
  "email": "Email address - MUST extract if present (look for @ symbol)",
  "phone": "Phone number with country code - MUST extract if present",
  "linkedin": "Full LinkedIn URL if present",
  "skills": ["Extract ALL technical skills, tools, technologies, soft skills mentioned"],
  "experience": [
    {
      "role": "Job title (if multiple roles at same company, use most recent)",
      "company": "Company name - MUST be unique, no duplicates",
      "duration": "Full date range at company (e.g., Jan 2020 - Present)",
      "achievements": ["Key achievement or responsibility with metrics if available"]
    }
  ],
  "education": [
    {
      "degree": "Degree name (e.g., B.Tech, MBA, M.Sc)",
      "institution": "University/College name",
      "year": "Graduation year or date range"
    }
  ],
  "summary": "Write a compelling 2-3 sentence professional summary based on the resume"
}

RULES:
1. Extract ALL skills - technical, tools, frameworks, languages, soft skills
2. Phone numbers: Include country code, common formats: +91-XXXXX, (123) 456-7890
3. Email: Look carefully throughout the document for @ symbol
4. Experience: List ALL positions, most recent first
5. Use empty string "" for missing text fields, empty array [] for missing lists
6. DO NOT use null values
7. Return ONLY the JSON, no markdown, no explanations"""

        # Use more CV text for better extraction (increased to 6000 chars)
        cv_text_to_use = cv_text[:6000] if len(cv_text) > 6000 else cv_text
        
        # Pre-extract contact info using regex as backup
        email_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cv_text)
        phone_match = re.search(r'(?:\+91[-\s]?)?(?:\d{10}|\d{5}[-\s]?\d{5}|\(\d{3}\)\s?\d{3}[-\s]?\d{4})', cv_text)
        linkedin_match = re.search(r'linkedin\.com/in/[a-zA-Z0-9-]+', cv_text)
        
        backup_email = email_match.group() if email_match else ""
        backup_phone = phone_match.group() if phone_match else ""
        backup_linkedin = f"https://{linkedin_match.group()}" if linkedin_match else ""
        
        prompt = f"""Extract ALL information from this resume. Pay special attention to contact details.

PRE-DETECTED CONTACT INFO (verify and include if correct):
- Email found: {backup_email}
- Phone found: {backup_phone}
- LinkedIn found: {backup_linkedin}

RESUME TEXT:
---
{cv_text_to_use}
---

Parse the resume thoroughly and return ONLY valid JSON. Include the contact info above if it looks correct."""
        
        print(f"[DEBUG] Parsing CV with {len(cv_text)} chars")
        print(f"[DEBUG] Regex backup - Email: {backup_email}, Phone: {backup_phone}")
        
        # Use OpenAI SDK directly
        response = await call_openai_directly(system_prompt, prompt, llm_key)
        
        print(f"[DEBUG] AI Response for parsing: {response[:800]}")
        
        # Extract JSON from response
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            parsed_data = json.loads(json_match.group())
            print(f"[DEBUG] Parsed data keys: {list(parsed_data.keys())}")
            print(f"[DEBUG] Email: {parsed_data.get('email')}, Phone: {parsed_data.get('phone')}")
            print(f"[DEBUG] Skills count: {len(parsed_data.get('skills', []))}")
            
            # Handle null values
            for key in ['name', 'current_role', 'email', 'phone', 'linkedin', 'summary']:
                if parsed_data.get(key) is None:
                    parsed_data[key] = "" if key != 'name' else "Candidate"
            
            # Use regex backup for contact info if AI missed it
            if not parsed_data.get('email') and backup_email:
                parsed_data['email'] = backup_email
                print(f"[DEBUG] Using regex backup email: {backup_email}")
            if not parsed_data.get('phone') and backup_phone:
                parsed_data['phone'] = backup_phone
                print(f"[DEBUG] Using regex backup phone: {backup_phone}")
            if not parsed_data.get('linkedin') and backup_linkedin:
                parsed_data['linkedin'] = backup_linkedin
                print(f"[DEBUG] Using regex backup linkedin: {backup_linkedin}")
            
            # Ensure lists are not None
            for key in ['skills', 'experience', 'education']:
                if parsed_data.get(key) is None:
                    parsed_data[key] = []
            
            # Deduplicate experience entries by company name
            if parsed_data.get('experience'):
                seen_companies = {}
                deduped_experience = []
                for exp in parsed_data['experience']:
                    company = exp.get('company', '').lower().strip()
                    if company and company not in seen_companies:
                        seen_companies[company] = True
                        deduped_experience.append(exp)
                    elif company:
                        print(f"[DEBUG] Deduped duplicate company: {exp.get('company')}")
                parsed_data['experience'] = deduped_experience
            
            return ParsedResume(**parsed_data)
        else:
            raise ValueError("No JSON found in response")
    except Exception as e:
        print(f"[ERROR] AI parsing error: {e}")
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        # Return existing data or minimal fallback
        if existing_data:
            return ParsedResume(**existing_data)
        return ParsedResume(
            name="CV Upload",
            summary="AI parsing failed - please edit manually"
        )

async def generate_candidate_story(candidate_data: dict, job_data: dict) -> CandidateStory:
    """Generate AI candidate story using RecruitAssist AI with accurate scoring"""
    llm_key = os.environ.get('EMERGENT_LLM_KEY')
    if not llm_key:
        # Return fallback story if no LLM key
        return CandidateStory(
            headline=f"Candidate for {job_data.get('title', 'Position')}",
            summary="AI story generation unavailable - LLM key not configured",
            timeline=[],
            skills=candidate_data.get('skills', []),
            fit_score=50,
            highlights=["Manual review recommended"]
        )
    
    try:
        # Enhanced RecruitAssist AI System Prompt for Story Generation
        system_prompt = """You are an expert recruiter analyzing candidate-job fit. Generate an ACCURATE candidate story.

CRITICAL RULES - READ CAREFULLY:

1. NEVER FABRICATE "CAREER TRANSITIONS":
   - Do NOT say candidate is "transitioning to" or "looking to move into" a new field unless explicitly stated in their CV
   - Do NOT assume someone from Field A wants to work in Field B
   - If a Risk Analyst applies for QA role, report them as a RISK ANALYST, not "transitioning QA professional"
   - Describe what they ARE, not what you imagine they want to be

2. FIT_SCORE MUST BE REALISTIC:
   - MISMATCHED PROFILES: If candidate's background doesn't match job requirements, score MUST be 15-30%
   - PARTIALLY MATCHING: Some overlap but different core domain = 30-50%
   - MODERATE MATCH: Same domain, some skill gaps = 50-70%
   - STRONG MATCH: Direct experience in role = 70-85%
   - EXCELLENT MATCH: Perfect fit = 85-95%
   
   Example: Risk Analyst applying for QA Tester = 20-30% (completely different domain)
   Example: Java Developer applying for Python Developer = 50-65% (same domain, different stack)
   Example: Senior QA applying for QA Lead = 75-85% (direct progression)

3. TIMELINE - USE ONLY CV DATA:
   - Extract actual job titles, companies, dates from the CV
   - NO invented companies, NO made-up achievements
   - If achievement not in CV, don't include it
   - Each unique company-role combination appears once

4. HEADLINE AND SUMMARY:
   - Describe candidate's ACTUAL background
   - DO NOT spin their experience to match the job
   - Be honest about domain mismatch if exists

Return ONLY this JSON structure:

{
  "headline": "One sentence describing candidate's ACTUAL expertise and background",
  "summary": "3-4 sentences about their REAL career journey. If mismatched for the role, state that honestly.",
  "timeline": [
    {"year": "2022-Present", "title": "Actual Job Title", "company": "Actual Company", "achievement": "Actual achievement from CV or 'Key responsibilities: [list from CV]'"}
  ],
  "skills": ["Actual skills from CV - top 10-15"],
  "highlights": [
    "Real achievement from CV",
    "Another real accomplishment"
  ],
  "fit_score": 25
}

IMPORTANT: If candidate is from a different domain than the job, the fit_score should be LOW (15-35). Do NOT try to make them seem like a fit."""

        # Build comprehensive candidate data
        experience_list = candidate_data.get('experience', [])
        essential_candidate_data = {
            "name": candidate_data.get('name', ''),
            "current_role": candidate_data.get('current_role', ''),
            "skills": candidate_data.get('skills', []),
            "summary": candidate_data.get('summary', ''),
            "experience": [
                {
                    "role": exp.get('role', ''),
                    "company": exp.get('company', ''),
                    "duration": exp.get('duration', ''),
                    "achievements": exp.get('achievements', [])
                }
                for exp in experience_list[:7]
            ],
            "education": candidate_data.get('education', [])[:5]
        }
        
        # Get job requirements
        job_skills = job_data.get('required_skills', [])
        exp_range = job_data.get('experience_range', {})
        
        # Determine job domain keywords
        job_title = job_data.get('title', 'Position').lower()
        job_domain_keywords = []
        if any(x in job_title for x in ['qa', 'test', 'quality']):
            job_domain_keywords = ['testing', 'qa', 'test automation', 'selenium', 'manual testing', 'bug', 'defect']
        elif any(x in job_title for x in ['developer', 'engineer', 'programmer']):
            job_domain_keywords = ['development', 'coding', 'programming', 'software', 'api', 'backend', 'frontend']
        elif any(x in job_title for x in ['analyst', 'data']):
            job_domain_keywords = ['analysis', 'data', 'analytics', 'reporting', 'sql', 'excel']
        
        prompt = f'''Analyze this candidate for the job and generate an HONEST story.

CANDIDATE DATA:
{json.dumps(essential_candidate_data, indent=2)}

JOB REQUIREMENTS:
- Title: {job_data.get('title', 'Position')}
- Description: {job_data.get('description', '')[:1500]}
- Required Skills: {', '.join(job_skills) if job_skills else 'Not specified'}
- Experience Required: {exp_range.get('min_years', 0)}-{exp_range.get('max_years', 10)} years
- Domain Keywords: {', '.join(job_domain_keywords)}

SCORING CHECKLIST:
1. Is candidate's current role in the same domain as "{job_data.get('title', '')}"? 
   - If NO (different domain): Start with base score of 20-30
   - If YES (same domain): Start with base score of 50-60
2. How many required skills does candidate have? Add 1-2 points per matching skill
3. Does experience level match? Adjust +/- 5-10 points

CANDIDATE'S DOMAIN: Based on their current role "{candidate_data.get('current_role', '')}", what domain are they in?
JOB'S DOMAIN: Based on title "{job_data.get('title', '')}", what domain is this job?

If domains are DIFFERENT (e.g., Risk/Analytics vs QA Testing), fit_score should be 15-35%.
Do NOT pretend they are transitioning or a good fit if they are not.

Generate ACCURATE JSON response.'''
        
        # Use OpenAI SDK directly
        response = await call_openai_directly(system_prompt, prompt, llm_key)
        
        print(f"[DEBUG] AI Story Response: {response[:1000]}")
        
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            story_data = json.loads(json_match.group())
            print(f"[DEBUG] Story fit_score from AI: {story_data.get('fit_score')}")
            print(f"[DEBUG] Timeline entries: {len(story_data.get('timeline', []))}")
            
            # Deduplicate timeline entries by company name
            if story_data.get('timeline'):
                seen_companies = {}
                deduped_timeline = []
                for entry in story_data['timeline']:
                    company = entry.get('company', '').lower().strip()
                    if company and company not in seen_companies:
                        seen_companies[company] = True
                        deduped_timeline.append(entry)
                    elif company:
                        print(f"[DEBUG] Deduped duplicate timeline company: {entry.get('company')}")
                story_data['timeline'] = deduped_timeline
            
            # Validate fit_score - only override if clearly wrong
            ai_fit_score = story_data.get('fit_score')
            if ai_fit_score is None or ai_fit_score == 0:
                # Calculate our own fit score
                print("[DEBUG] AI didn't provide fit_score, calculating...")
                story_data['fit_score'] = calculate_fit_score(candidate_data, job_data)
            
            # Ensure all fields have values
            if not story_data.get('headline'):
                story_data['headline'] = f"{candidate_data.get('name', 'Candidate')} - {candidate_data.get('current_role', 'Professional')}"
            if not story_data.get('summary'):
                story_data['summary'] = candidate_data.get('summary', 'Professional candidate profile')
            if not story_data.get('highlights'):
                story_data['highlights'] = []
            if not story_data.get('timeline'):
                # Build timeline from experience if AI didn't provide it
                seen_companies = {}
                deduped_exp_timeline = []
                for exp in candidate_data.get('experience', [])[:5]:
                    company = exp.get('company', '').lower().strip()
                    if company and company not in seen_companies:
                        seen_companies[company] = True
                        deduped_exp_timeline.append({
                            "year": exp.get('duration', ''),
                            "title": exp.get('role', ''),
                            "company": exp.get('company', ''),
                            "achievement": exp.get('achievements', [''])[0] if exp.get('achievements') else ''
                        })
                story_data['timeline'] = deduped_exp_timeline
            if not story_data.get('skills'):
                story_data['skills'] = candidate_data.get('skills', [])[:15]
                
            return CandidateStory(**story_data)
        else:
            raise ValueError("No JSON found in response")
    except Exception as e:
        print(f"[ERROR] AI story generation error: {e}")
        import traceback
        print(f"[ERROR] Traceback: {traceback.format_exc()}")
        # Return calculated story if AI fails
        fit_score = calculate_fit_score(candidate_data, job_data)
        return CandidateStory(
            headline=f"{candidate_data.get('name', 'Candidate')} for {job_data.get('title', 'Position')}",
            summary=candidate_data.get('summary', 'Professional candidate - AI story generation failed'),
            timeline=[
                {
                    "year": exp.get('duration', ''),
                    "title": exp.get('role', ''),
                    "company": exp.get('company', ''),
                    "achievement": exp.get('achievements', [''])[0] if exp.get('achievements') else ''
                }
                for exp in candidate_data.get('experience', [])[:5]
            ],
            skills=candidate_data.get('skills', [])[:15],
            fit_score=fit_score,
            highlights=["Review candidate profile for details"]
        )


def calculate_fit_score(candidate_data: dict, job_data: dict) -> int:
    """Calculate fit score based on skills, experience, and role alignment"""
    
    # Skills Match (45% weight)
    candidate_skills = set([s.lower().strip() for s in candidate_data.get('skills', [])])
    job_skills = set([s.lower().strip() for s in job_data.get('required_skills', [])])
    
    if job_skills:
        # Direct match
        direct_matches = len(candidate_skills & job_skills)
        # Partial match (check if job skill is substring of candidate skill or vice versa)
        partial_matches = 0
        for js in job_skills:
            if js not in candidate_skills:
                for cs in candidate_skills:
                    if js in cs or cs in js:
                        partial_matches += 0.5
                        break
        total_matches = direct_matches + partial_matches
        skills_match_score = min((total_matches / len(job_skills)) * 45, 45)
    else:
        skills_match_score = 22  # Base score if no job skills specified
    
    # Experience Match (35% weight)
    experience_list = candidate_data.get('experience', [])
    candidate_years = len(experience_list) * 2  # Rough estimate: 2 years per position
    
    # Try to extract actual years from duration strings
    total_months = 0
    for exp in experience_list:
        duration = exp.get('duration', '').lower()
        if 'present' in duration or 'current' in duration:
            total_months += 24  # Assume 2 years if current
        elif '-' in duration:
            try:
                parts = duration.split('-')
                if len(parts) == 2:
                    # Try to parse years
                    import re
                    years = re.findall(r'20\d{2}|19\d{2}', duration)
                    if len(years) >= 2:
                        total_months += (int(years[1]) - int(years[0])) * 12
                    else:
                        total_months += 24
            except:
                total_months += 24
        else:
            total_months += 24
    
    candidate_years = max(candidate_years, total_months // 12)
    
    exp_range = job_data.get('experience_range', {})
    min_years = exp_range.get('min_years', 0)
    max_years = exp_range.get('max_years', 15)
    
    if min_years <= candidate_years <= max_years:
        exp_match_score = 35
    elif candidate_years > max_years:
        exp_match_score = 30  # Overqualified but still good
    elif candidate_years >= min_years * 0.7:
        exp_match_score = 25  # Close enough
    else:
        exp_match_score = max((candidate_years / min_years) * 35, 10) if min_years > 0 else 20
    
    # Role Alignment (20% weight)
    candidate_role = (candidate_data.get('current_role') or "").lower()
    job_title = (job_data.get('title') or "").lower()
    
    # Check for keyword matches
    job_keywords = set(job_title.replace('-', ' ').replace('/', ' ').split())
    role_keywords = set(candidate_role.replace('-', ' ').replace('/', ' ').split())
    
    common_keywords = job_keywords & role_keywords
    if common_keywords:
        role_match_score = min(len(common_keywords) * 5 + 10, 20)
    else:
        # Check for related terms
        related_terms = {
            'developer': ['engineer', 'programmer', 'coder'],
            'engineer': ['developer', 'architect', 'designer'],
            'manager': ['lead', 'head', 'director', 'supervisor'],
            'analyst': ['consultant', 'specialist', 'advisor'],
            'designer': ['ux', 'ui', 'creative', 'artist']
        }
        role_match_score = 8
        for key, synonyms in related_terms.items():
            if key in job_title:
                for syn in synonyms:
                    if syn in candidate_role:
                        role_match_score = 15
                        break
    
    total_score = int(skills_match_score + exp_match_score + role_match_score)
    final_score = min(max(total_score, 20), 100)
    
    print(f"[DEBUG] Fit Score Calculation: Skills={skills_match_score:.1f}, Exp={exp_match_score:.1f}, Role={role_match_score} = {final_score}%")
    
    return final_score

def create_access_token(data: dict) -> str:
    """Create a JWT access token"""
    to_encode = data.copy()
    expire = datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRATION_HOURS)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)
    return encoded_jwt

def decode_token(token: str) -> dict:
    """Decode and verify JWT token"""
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Token has expired"
        )
    except jwt.InvalidTokenError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token"
        )

async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> dict:
    """Dependency to get current authenticated user"""
    token = credentials.credentials
    payload = decode_token(token)
    
    email = payload.get("email")
    if not email:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication credentials"
        )
    
    user = await db.users.find_one({"email": email}, {"_id": 0, "password_hash": 0})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="User not found"
        )
    
    return user

async def require_role(required_roles: list[str]):
    """Dependency factory to require specific roles"""
    async def role_checker(current_user: dict = Depends(get_current_user)):
        if current_user["role"] not in required_roles:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Insufficient permissions"
            )
        return current_user
    return role_checker


# ============ GOVERNANCE HELPERS ============

async def log_audit_event(
    user_id: str,
    user_email: str,
    user_role: str,
    action_type: str,
    entity_type: str,
    entity_id: Optional[str] = None,
    client_id: Optional[str] = None,
    previous_value: Optional[dict] = None,
    new_value: Optional[dict] = None,
    metadata: Optional[dict] = None,
    ip_address: Optional[str] = None
):
    """Log an audit event to the audit_logs collection"""
    log_entry = {
        "log_id": f"log_{uuid.uuid4().hex[:12]}",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "user_id": user_id,
        "user_email": user_email,
        "user_role": user_role,
        "client_id": client_id,
        "action_type": action_type,
        "entity_type": entity_type,
        "entity_id": entity_id,
        "previous_value": previous_value,
        "new_value": new_value,
        "metadata": metadata or {},
        "ip_address": ip_address
    }
    
    await db.audit_logs.insert_one(log_entry)
    print(f"[AUDIT] {action_type} by {user_email} on {entity_type} {entity_id}")

async def get_user_permissions(user: dict, client_id: Optional[str] = None) -> PermissionSet:
    """Get aggregated permissions for a user in a specific client context"""
    
    # Arbeit Admin bypass - full permissions
    if user["role"] in ["admin", "recruiter"]:
        return PermissionSet(
            can_view_jobs=True,
            can_create_jobs=True,
            can_edit_jobs=True,
            can_delete_jobs=True,
            can_view_candidates=True,
            can_create_candidates=True,
            can_edit_candidates=True,
            can_delete_candidates=True,
            can_update_candidate_status=True,
            can_upload_cv=True,
            can_replace_cv=True,
            can_regenerate_story=True,
            can_view_full_cv=True,
            can_view_redacted_cv=True,
            can_view_audit_log=True,
            can_manage_roles=True,
            can_manage_users=True,
            can_export_reports=True
        )
    
    # For client_user, get their assigned roles
    if not client_id:
        client_id = user.get("client_id")
    
    if not client_id:
        # No client context, return minimal permissions
        return PermissionSet()
    
    # Get user's role assignments for this client
    role_assignments = await db.user_client_roles.find({
        "user_id": user.get("user_id", user.get("email")),
        "client_id": client_id
    }).to_list(100)
    
    if not role_assignments:
        # No roles assigned, give client users basic operational permissions
        return PermissionSet(
            can_view_jobs=True,
            can_create_jobs=True,
            can_edit_jobs=True,
            can_view_candidates=True,
            can_create_candidates=True,
            can_edit_candidates=True,
            can_update_candidate_status=True,
            can_upload_cv=True,
            can_view_redacted_cv=True
        )
    
    # Aggregate permissions from all assigned roles
    aggregated_perms = {}
    for assignment in role_assignments:
        role = await db.client_roles.find_one({"role_id": assignment["client_role_id"]}, {"_id": 0})
        if role and "permissions" in role:
            for key, value in role["permissions"].items():
                # OR logic: if any role grants permission, user has it
                if value is True:
                    aggregated_perms[key] = True
    
    # Create PermissionSet with aggregated permissions
    return PermissionSet(**aggregated_perms) if aggregated_perms else PermissionSet()

async def check_permission(user: dict, permission: str, client_id: Optional[str] = None) -> bool:
    """Check if user has a specific permission"""
    perms = await get_user_permissions(user, client_id)
    return getattr(perms, permission, False)

def requires_permission(permission: str):
    """Decorator to enforce permission check on endpoints"""
    async def permission_checker(
        client_id: Optional[str] = None,
        current_user: dict = Depends(get_current_user)
    ):
        # Admin bypass
        if current_user["role"] in ["admin", "recruiter"]:
            return current_user
        
        # Check permission
        has_permission = await check_permission(current_user, permission, client_id)
        
        if not has_permission:
            # Log failed access attempt
            await log_audit_event(
                user_id=current_user.get("user_id", current_user["email"]),
                user_email=current_user["email"],
                user_role=current_user["role"],
                action_type="ACCESS_DENIED",
                entity_type="permission",
                client_id=client_id,
                metadata={"required_permission": permission}
            )
            
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"Permission denied: {permission} required"
            )
        
        return current_user
    
    return permission_checker

# Default role templates
DEFAULT_ROLE_TEMPLATES = {
    "Client Owner": {
        "description": "Full access to manage jobs, candidates, and view reports",
        "permissions": PermissionSet(
            can_view_jobs=True,
            can_create_jobs=True,
            can_edit_jobs=True,
            can_delete_jobs=False,  # Only admin can delete
            can_view_candidates=True,
            can_create_candidates=True,
            can_edit_candidates=True,
            can_delete_candidates=False,
            can_update_candidate_status=True,
            can_upload_cv=True,
            can_replace_cv=True,
            can_regenerate_story=True,
            can_view_full_cv=True,
            can_view_redacted_cv=True,
            can_view_audit_log=True,
            can_manage_roles=True,
            can_manage_users=True,
            can_export_reports=True
        )
    },
    "Hiring Manager": {
        "description": "Can manage jobs and candidates, upload CVs, and update status",
        "permissions": PermissionSet(
            can_view_jobs=True,
            can_create_jobs=True,
            can_edit_jobs=True,
            can_delete_jobs=False,
            can_view_candidates=True,
            can_create_candidates=True,
            can_edit_candidates=True,
            can_delete_candidates=False,
            can_update_candidate_status=True,
            can_upload_cv=True,
            can_replace_cv=True,
            can_regenerate_story=False,
            can_view_full_cv=False,
            can_view_redacted_cv=True,
            can_view_audit_log=False,
            can_manage_roles=False,
            can_manage_users=False,
            can_export_reports=False
        )
    },
    "Interviewer": {
        "description": "Read-only access to view jobs and candidates",
        "permissions": PermissionSet(
            can_view_jobs=True,
            can_create_jobs=False,
            can_edit_jobs=False,
            can_delete_jobs=False,
            can_view_candidates=True,
            can_create_candidates=False,
            can_edit_candidates=False,
            can_delete_candidates=False,
            can_update_candidate_status=False,
            can_upload_cv=False,
            can_replace_cv=False,
            can_regenerate_story=False,
            can_view_full_cv=False,
            can_view_redacted_cv=True,
            can_view_audit_log=False,
            can_manage_roles=False,
            can_manage_users=False,
            can_export_reports=False
        )
    }
}

async def create_default_roles_for_client(client_id: str):
    """Create default role templates when a new client is created"""
    for role_name, role_config in DEFAULT_ROLE_TEMPLATES.items():
        role_doc = {
            "role_id": f"role_{uuid.uuid4().hex[:12]}",
            "client_id": client_id,
            "name": role_name,
            "description": role_config["description"],
            "permissions": role_config["permissions"].model_dump(),
            "created_at": datetime.now(timezone.utc).isoformat(),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }
        await db.client_roles.insert_one(role_doc)
        print(f"[RBAC] Created default role '{role_name}' for client {client_id}")


# ============ AUTH ROUTES ============

@api_router.post("/auth/register", response_model=UserResponse)
async def register_user(user_data: UserCreate):
    """Register a new user (typically done by admin)"""
    
    # Check if user already exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Email already registered"
        )
    
    # Validate client_id for client_user role
    if user_data.role == "client_user":
        if not user_data.client_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="client_id is required for client_user role"
            )
        # Verify client exists
        client = await db.clients.find_one({"client_id": user_data.client_id})
        if not client:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid client_id"
            )
    
    # Hash password and create user
    password_hash = hash_password(user_data.password)
    
    user_doc = {
        "email": user_data.email,
        "name": user_data.name,
        "role": user_data.role,
        "client_id": user_data.client_id,
        "password_hash": password_hash,
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.users.insert_one(user_doc)
    
    return UserResponse(
        email=user_data.email,
        name=user_data.name,
        role=user_data.role,
        client_id=user_data.client_id
    )

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(credentials: UserLogin):
    """Login and receive JWT token"""
    
    # Find user by email
    user = await db.users.find_one({"email": credentials.email}, {"_id": 0})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    # Verify password
    if not verify_password(credentials.password, user["password_hash"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    # Create access token
    token_data = {
        "email": user["email"],
        "role": user["role"],
        "client_id": user.get("client_id")
    }
    access_token = create_access_token(token_data)
    
    return TokenResponse(
        access_token=access_token,
        token_type="bearer",
        user=UserResponse(
            email=user["email"],
            name=user["name"],
            role=user["role"],
            client_id=user.get("client_id")
        ),
        must_change_password=user.get("must_change_password", False)
    )

@api_router.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: dict = Depends(get_current_user)):
    """Get current authenticated user"""
    return UserResponse(
        email=current_user["email"],
        name=current_user["name"],
        role=current_user["role"],
        client_id=current_user.get("client_id")
    )


class ChangePasswordRequest(BaseModel):
    current_password: str
    new_password: str


@api_router.post("/auth/change-password")
async def change_password(
    request: ChangePasswordRequest,
    current_user: dict = Depends(get_current_user)
):
    """Change password for the current user (works for all user types including client_user)"""
    # Get the full user document
    user = await db.users.find_one({"email": current_user["email"]})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    # Verify current password
    if not verify_password(request.current_password, user["password_hash"]):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Current password is incorrect"
        )
    
    # Hash the new password
    new_password_hash = hash_password(request.new_password)
    
    # Update password and clear must_change_password flag
    await db.users.update_one(
        {"email": current_user["email"]},
        {
            "$set": {
                "password_hash": new_password_hash,
                "must_change_password": False
            }
        }
    )
    
    return {"message": "Password changed successfully"}


# ============ CANDIDATE PORTAL AUTHENTICATION ============

@api_router.post("/candidate-portal/register", response_model=CandidatePortalResponse)
async def register_candidate_portal(candidate_data: CandidatePortalRegister):
    """Register a new candidate for the portal"""
    # Check if candidate already exists
    existing = await db.candidate_portal_users.find_one({"email": candidate_data.email})
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="A candidate with this email already exists"
        )
    
    # Hash password
    password_hash = bcrypt.hashpw(candidate_data.password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    # Generate candidate portal ID
    candidate_portal_id = f"cp_{uuid.uuid4().hex[:12]}"
    
    candidate_doc = {
        "candidate_portal_id": candidate_portal_id,
        "email": candidate_data.email,
        "name": candidate_data.name,
        "phone": candidate_data.phone,
        "linkedin_url": candidate_data.linkedin_url,
        "current_company": candidate_data.current_company,
        "experience_years": candidate_data.experience_years,
        "password_hash": password_hash,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "is_active": True
    }
    
    await db.candidate_portal_users.insert_one(candidate_doc)
    
    # Try to link with existing candidate records by email
    await db.candidates.update_many(
        {"email": candidate_data.email},
        {"$set": {"candidate_portal_id": candidate_portal_id}}
    )
    
    return CandidatePortalResponse(
        candidate_portal_id=candidate_portal_id,
        email=candidate_data.email,
        name=candidate_data.name,
        phone=candidate_data.phone,
        linkedin_url=candidate_data.linkedin_url,
        current_company=candidate_data.current_company,
        experience_years=candidate_data.experience_years,
        created_at=candidate_doc["created_at"]
    )


@api_router.post("/candidate-portal/login", response_model=CandidatePortalTokenResponse)
async def login_candidate_portal(login_data: CandidatePortalLogin):
    """Login for candidate portal"""
    candidate = await db.candidate_portal_users.find_one({"email": login_data.email})
    
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    if not bcrypt.checkpw(login_data.password.encode('utf-8'), candidate["password_hash"].encode('utf-8')):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid email or password"
        )
    
    # Generate JWT token
    token_data = {
        "email": candidate["email"],
        "candidate_portal_id": candidate["candidate_portal_id"],
        "type": "candidate_portal",
        "exp": datetime.now(timezone.utc) + timedelta(hours=24)
    }
    
    token = jwt.encode(token_data, JWT_SECRET, algorithm="HS256")
    
    must_change = candidate.get("must_change_password", False)
    
    return CandidatePortalTokenResponse(
        access_token=token,
        token_type="bearer",
        candidate=CandidatePortalResponse(
            candidate_portal_id=candidate["candidate_portal_id"],
            email=candidate["email"],
            name=candidate["name"],
            phone=candidate["phone"],
            linkedin_url=candidate.get("linkedin_url"),
            current_company=candidate.get("current_company"),
            experience_years=candidate.get("experience_years"),
            created_at=candidate["created_at"],
            must_change_password=must_change
        ),
        must_change_password=must_change
    )


async def get_current_candidate(credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer())):
    """Dependency to get current authenticated candidate"""
    try:
        token = credentials.credentials
        payload = jwt.decode(token, JWT_SECRET, algorithms=["HS256"])
        
        if payload.get("type") != "candidate_portal":
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid token type"
            )
        
        candidate = await db.candidate_portal_users.find_one(
            {"candidate_portal_id": payload["candidate_portal_id"]},
            {"_id": 0, "password_hash": 0}
        )
        
        if not candidate:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Candidate not found"
            )
        
        return candidate
    except jwt.ExpiredSignatureError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Token has expired"
        )
    except jwt.InvalidTokenError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token"
        )


@api_router.get("/candidate-portal/me", response_model=CandidatePortalResponse)
async def get_candidate_me(current_candidate: dict = Depends(get_current_candidate)):
    """Get current authenticated candidate"""
    return CandidatePortalResponse(
        candidate_portal_id=current_candidate["candidate_portal_id"],
        email=current_candidate["email"],
        name=current_candidate["name"],
        phone=current_candidate["phone"],
        linkedin_url=current_candidate.get("linkedin_url"),
        current_company=current_candidate.get("current_company"),
        experience_years=current_candidate.get("experience_years"),
        created_at=current_candidate["created_at"],
        must_change_password=current_candidate.get("must_change_password", False)
    )


@api_router.post("/candidate-portal/change-password")
async def change_candidate_password(
    password_data: CandidatePasswordChange,
    current_candidate: dict = Depends(get_current_candidate)
):
    """Change password for candidate portal user"""
    # Fetch full user data including password_hash
    full_user = await db.candidate_portal_users.find_one(
        {"candidate_portal_id": current_candidate["candidate_portal_id"]},
        {"_id": 0}
    )
    
    if not full_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    # Verify current password
    if not bcrypt.checkpw(password_data.current_password.encode('utf-8'), 
                          full_user["password_hash"].encode('utf-8')):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Current password is incorrect"
        )
    
    # Validate new password
    if len(password_data.new_password) < 6:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="New password must be at least 6 characters"
        )
    
    # Hash new password
    new_hash = bcrypt.hashpw(password_data.new_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    # Update password and remove must_change_password flag
    await db.candidate_portal_users.update_one(
        {"candidate_portal_id": current_candidate["candidate_portal_id"]},
        {"$set": {"password_hash": new_hash, "must_change_password": False}}
    )
    
    return {"message": "Password changed successfully"}


# ============ ADMIN: CANDIDATE PORTAL MANAGEMENT ============

class CandidatePortalAdminResponse(BaseModel):
    """Admin view of candidate portal user"""
    candidate_portal_id: str
    email: str
    name: str
    phone: str
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    experience_years: Optional[int] = None
    created_at: str
    must_change_password: bool = False
    linked_candidate_id: Optional[str] = None
    status: Optional[str] = "active"


class CandidatePortalCreateByAdmin(BaseModel):
    """Admin create candidate portal user"""
    email: EmailStr
    name: str
    phone: str
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    experience_years: Optional[int] = None
    link_to_candidate_id: Optional[str] = None
    send_welcome_email: bool = True


class CandidatePortalUpdateByAdmin(BaseModel):
    """Admin update candidate portal user"""
    name: Optional[str] = None
    phone: Optional[str] = None
    linkedin_url: Optional[str] = None
    current_company: Optional[str] = None
    experience_years: Optional[int] = None
    status: Optional[str] = None


@api_router.get("/admin/candidate-portal-users", response_model=list[CandidatePortalAdminResponse])
async def list_candidate_portal_users(
    search: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """List all candidate portal users (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    query = {}
    
    if search:
        query["$or"] = [
            {"name": {"$regex": search, "$options": "i"}},
            {"email": {"$regex": search, "$options": "i"}},
            {"phone": {"$regex": search, "$options": "i"}}
        ]
    
    if status and status != "all":
        query["status"] = status
    
    users = await db.candidate_portal_users.find(
        query,
        {"_id": 0, "password_hash": 0}
    ).sort("created_at", -1).to_list(500)
    
    return [CandidatePortalAdminResponse(
        candidate_portal_id=u["candidate_portal_id"],
        email=u["email"],
        name=u["name"],
        phone=u.get("phone", ""),
        linkedin_url=u.get("linkedin_url"),
        current_company=u.get("current_company"),
        experience_years=u.get("experience_years"),
        created_at=u["created_at"],
        must_change_password=u.get("must_change_password", False),
        linked_candidate_id=u.get("linked_candidate_id"),
        status=u.get("status", "active")
    ) for u in users]


@api_router.get("/admin/candidate-portal-users/{portal_id}", response_model=CandidatePortalAdminResponse)
async def get_candidate_portal_user(
    portal_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get a specific candidate portal user (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    user = await db.candidate_portal_users.find_one(
        {"candidate_portal_id": portal_id},
        {"_id": 0, "password_hash": 0}
    )
    
    if not user:
        raise HTTPException(status_code=404, detail="Candidate portal user not found")
    
    return CandidatePortalAdminResponse(
        candidate_portal_id=user["candidate_portal_id"],
        email=user["email"],
        name=user["name"],
        phone=user.get("phone", ""),
        linkedin_url=user.get("linkedin_url"),
        current_company=user.get("current_company"),
        experience_years=user.get("experience_years"),
        created_at=user["created_at"],
        must_change_password=user.get("must_change_password", False),
        linked_candidate_id=user.get("linked_candidate_id"),
        status=user.get("status", "active")
    )


@api_router.post("/admin/candidate-portal-users", response_model=CandidatePortalAdminResponse)
async def create_candidate_portal_user_by_admin(
    user_data: CandidatePortalCreateByAdmin,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Create a candidate portal user (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    from notification_service import send_email, get_candidate_selection_email_template
    import secrets
    
    # Check if email already exists
    existing = await db.candidate_portal_users.find_one({"email": user_data.email})
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="A candidate portal user with this email already exists"
        )
    
    # Generate temp password
    temp_password = secrets.token_urlsafe(8)
    password_hash = hash_password(temp_password)
    
    candidate_portal_id = f"cp_{uuid.uuid4().hex[:12]}"
    now = datetime.now(timezone.utc).isoformat()
    
    user_doc = {
        "candidate_portal_id": candidate_portal_id,
        "email": user_data.email,
        "name": user_data.name,
        "phone": user_data.phone,
        "linkedin_url": user_data.linkedin_url,
        "current_company": user_data.current_company,
        "experience_years": user_data.experience_years,
        "password_hash": password_hash,
        "must_change_password": True,
        "status": "active",
        "linked_candidate_id": user_data.link_to_candidate_id,
        "created_at": now,
        "created_by": current_user["email"]
    }
    
    await db.candidate_portal_users.insert_one(user_doc)
    
    # Link to candidate if specified
    if user_data.link_to_candidate_id:
        await db.candidates.update_one(
            {"candidate_id": user_data.link_to_candidate_id},
            {"$set": {"candidate_portal_id": candidate_portal_id}}
        )
    
    # Send welcome email
    if user_data.send_welcome_email:
        frontend_url = os.environ.get('REACT_APP_FRONTEND_URL', '')
        
        # Create a simple welcome email
        subject = "Welcome to Arbeit Talent Portal - Your Account is Ready"
        body = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
                .header {{ background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%); color: white; padding: 20px; border-radius: 8px 8px 0 0; text-align: center; }}
                .content {{ background: #f8fafc; padding: 20px; border: 1px solid #e2e8f0; }}
                .credentials {{ background: #1e293b; color: white; padding: 15px; border-radius: 8px; margin: 15px 0; }}
                .cred-label {{ color: #94a3b8; font-size: 12px; }}
                .cred-value {{ font-family: monospace; background: #334155; padding: 8px; border-radius: 4px; margin: 5px 0 15px 0; }}
                .warning {{ background: #fef3c7; border-left: 4px solid #f59e0b; padding: 10px; margin: 15px 0; }}
                .btn {{ background: #3b82f6; color: white; padding: 12px 24px; text-decoration: none; border-radius: 6px; display: inline-block; }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>Welcome to Arbeit Talent Portal</h1>
                </div>
                <div class="content">
                    <p>Dear {user_data.name},</p>
                    <p>Your candidate portal account has been created. You can now log in to view and manage your interview schedules.</p>
                    
                    <div class="credentials">
                        <div class="cred-label">Portal URL</div>
                        <div class="cred-value">{frontend_url}/candidate/login</div>
                        <div class="cred-label">Email / Username</div>
                        <div class="cred-value">{user_data.email}</div>
                        <div class="cred-label">Temporary Password</div>
                        <div class="cred-value">{temp_password}</div>
                    </div>
                    
                    <div class="warning">
                        <strong>Important:</strong> You will be required to change your password on first login.
                    </div>
                    
                    <p style="text-align: center;">
                        <a href="{frontend_url}/candidate/login" class="btn">Login to Portal</a>
                    </p>
                </div>
            </div>
        </body>
        </html>
        """
        
        background_tasks.add_task(send_email, user_data.email, subject, body)
    
    return CandidatePortalAdminResponse(
        candidate_portal_id=candidate_portal_id,
        email=user_data.email,
        name=user_data.name,
        phone=user_data.phone,
        linkedin_url=user_data.linkedin_url,
        current_company=user_data.current_company,
        experience_years=user_data.experience_years,
        created_at=now,
        must_change_password=True,
        linked_candidate_id=user_data.link_to_candidate_id,
        status="active"
    )


@api_router.put("/admin/candidate-portal-users/{portal_id}", response_model=CandidatePortalAdminResponse)
async def update_candidate_portal_user_by_admin(
    portal_id: str,
    user_data: CandidatePortalUpdateByAdmin,
    current_user: dict = Depends(get_current_user)
):
    """Update a candidate portal user (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    user = await db.candidate_portal_users.find_one({"candidate_portal_id": portal_id})
    if not user:
        raise HTTPException(status_code=404, detail="Candidate portal user not found")
    
    update_data = user_data.model_dump(exclude_unset=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="No update data provided")
    
    await db.candidate_portal_users.update_one(
        {"candidate_portal_id": portal_id},
        {"$set": update_data}
    )
    
    updated = await db.candidate_portal_users.find_one(
        {"candidate_portal_id": portal_id},
        {"_id": 0, "password_hash": 0}
    )
    
    return CandidatePortalAdminResponse(
        candidate_portal_id=updated["candidate_portal_id"],
        email=updated["email"],
        name=updated["name"],
        phone=updated.get("phone", ""),
        linkedin_url=updated.get("linkedin_url"),
        current_company=updated.get("current_company"),
        experience_years=updated.get("experience_years"),
        created_at=updated["created_at"],
        must_change_password=updated.get("must_change_password", False),
        linked_candidate_id=updated.get("linked_candidate_id"),
        status=updated.get("status", "active")
    )


@api_router.delete("/admin/candidate-portal-users/{portal_id}")
async def delete_candidate_portal_user(
    portal_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a candidate portal user (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    user = await db.candidate_portal_users.find_one({"candidate_portal_id": portal_id})
    if not user:
        raise HTTPException(status_code=404, detail="Candidate portal user not found")
    
    # Remove link from candidate if exists
    await db.candidates.update_many(
        {"candidate_portal_id": portal_id},
        {"$unset": {"candidate_portal_id": ""}}
    )
    
    # Delete the portal user
    await db.candidate_portal_users.delete_one({"candidate_portal_id": portal_id})
    
    return {"message": f"Candidate portal user {user['email']} deleted successfully"}


@api_router.post("/admin/candidate-portal-users/{portal_id}/reset-password")
async def reset_candidate_portal_password(
    portal_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Reset password for a candidate portal user and send email (Admin/Recruiter only)"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(status_code=403, detail="Admin or Recruiter access required")
    
    from notification_service import send_email
    import secrets
    
    user = await db.candidate_portal_users.find_one({"candidate_portal_id": portal_id})
    if not user:
        raise HTTPException(status_code=404, detail="Candidate portal user not found")
    
    # Generate new temp password
    temp_password = secrets.token_urlsafe(8)
    password_hash = hash_password(temp_password)
    
    await db.candidate_portal_users.update_one(
        {"candidate_portal_id": portal_id},
        {"$set": {"password_hash": password_hash, "must_change_password": True}}
    )
    
    # Send email with new password
    frontend_url = os.environ.get('REACT_APP_FRONTEND_URL', '')
    subject = "Arbeit Talent Portal - Password Reset"
    body = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
            .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
            .header {{ background: #1e3a8a; color: white; padding: 20px; text-align: center; }}
            .content {{ background: #f8fafc; padding: 20px; }}
            .password {{ background: #1e293b; color: white; padding: 15px; font-family: monospace; border-radius: 8px; text-align: center; font-size: 18px; margin: 15px 0; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header"><h2>Password Reset</h2></div>
            <div class="content">
                <p>Dear {user['name']},</p>
                <p>Your password has been reset. Please use the new temporary password below to log in:</p>
                <div class="password">{temp_password}</div>
                <p>You will be required to change this password on your next login.</p>
                <p><a href="{frontend_url}/candidate/login">Click here to login</a></p>
            </div>
        </div>
    </body>
    </html>
    """
    
    background_tasks.add_task(send_email, user['email'], subject, body)
    
    return {"message": f"Password reset email sent to {user['email']}"}


@api_router.post("/candidates/{candidate_id}/send-selection-notification")
async def send_selection_notification(
    candidate_id: str,
    request: SendSelectionNotificationRequest = None,
    current_user: dict = Depends(get_current_user)
):
    """Send selection notification to candidate with portal login credentials"""
    from notification_service import send_email, get_candidate_selection_email_template
    import secrets
    
    # Check permissions - only admin and recruiters can send notifications
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admins and recruiters can send selection notifications"
        )
    
    # Get candidate
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    # Get candidate's email - try from candidate record or parsed data
    candidate_email = candidate.get("email")
    if not candidate_email:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Candidate does not have an email address. Please update the candidate's email first."
        )
    
    # Get job details
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    
    # Get client details
    client = await db.clients.find_one({"client_id": job["client_id"]}, {"_id": 0})
    if not client:
        client = {"company_name": "Unknown Company"}
    
    # Check if candidate already has portal account
    existing_portal_user = await db.candidate_portal_users.find_one({"email": candidate_email})
    
    if existing_portal_user:
        # Reset password for existing user
        temp_password = secrets.token_urlsafe(8)
        password_hash = bcrypt.hashpw(temp_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        
        await db.candidate_portal_users.update_one(
            {"email": candidate_email},
            {"$set": {
                "password_hash": password_hash,
                "must_change_password": True
            }}
        )
        
        candidate_portal_id = existing_portal_user["candidate_portal_id"]
    else:
        # Create new portal account
        temp_password = secrets.token_urlsafe(8)
        password_hash = bcrypt.hashpw(temp_password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
        
        candidate_portal_id = f"cp_{uuid.uuid4().hex[:12]}"
        portal_user_doc = {
            "candidate_portal_id": candidate_portal_id,
            "email": candidate_email,
            "name": candidate.get("name", "Candidate"),
            "phone": candidate.get("phone", ""),
            "linkedin_url": candidate.get("linkedin", ""),
            "current_company": candidate.get("current_company", ""),
            "experience_years": None,
            "password_hash": password_hash,
            "must_change_password": True,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "is_active": True
        }
        
        await db.candidate_portal_users.insert_one(portal_user_doc)
    
    # Link candidate record to portal user
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {"$set": {"candidate_portal_id": candidate_portal_id}}
    )
    
    # Update candidate status to SHORTLISTED
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {"$set": {"status": "SHORTLISTED"}}
    )
    
    # Get portal URL from environment
    portal_url = os.environ.get("FRONTEND_URL", "https://hirematch-52.preview.emergentagent.com")
    
    # Generate email
    subject, body = get_candidate_selection_email_template(
        candidate=candidate,
        job=job,
        client=client,
        login_email=candidate_email,
        temp_password=temp_password,
        portal_url=portal_url
    )
    
    # Send email
    email_result = await send_email(candidate_email, subject, body)
    
    # Create in-app notification record
    await db.notifications.insert_one({
        "notification_id": f"notif_{uuid.uuid4().hex[:8]}",
        "user_id": candidate_portal_id,
        "user_type": "candidate",
        "type": "SELECTION_NOTIFICATION",
        "title": "You've been selected!",
        "message": f"Congratulations! You've been selected for {job.get('title', 'a position')} at {client.get('company_name', 'our client')}",
        "is_read": False,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "metadata": {
            "job_id": job["job_id"],
            "candidate_id": candidate_id
        }
    })
    
    # Log audit
    await db.audit_logs.insert_one({
        "log_id": f"log_{uuid.uuid4().hex[:8]}",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "user_id": current_user.get("user_id", current_user["email"]),
        "user_email": current_user["email"],
        "action_type": "SELECTION_NOTIFICATION_SENT",
        "entity_type": "candidate",
        "entity_id": candidate_id,
        "metadata": {
            "candidate_email": candidate_email,
            "job_id": job["job_id"],
            "email_sent": email_result.get("success", False)
        }
    })
    
    return {
        "message": "Selection notification sent successfully",
        "email_sent": email_result.get("success", False),
        "candidate_email": candidate_email,
        "portal_account_created": not existing_portal_user
    }


@api_router.get("/candidate-portal/my-interviews")
async def get_candidate_interviews(current_candidate: dict = Depends(get_current_candidate)):
    """Get all interviews for the logged-in candidate"""
    # Find candidate records linked to this portal user
    candidate_records = await db.candidates.find(
        {"$or": [
            {"candidate_portal_id": current_candidate["candidate_portal_id"]},
            {"email": current_candidate["email"]}
        ]},
        {"_id": 0}
    ).to_list(100)
    
    candidate_ids = [c["candidate_id"] for c in candidate_records]
    
    if not candidate_ids:
        return []
    
    # Find all interviews for these candidates
    interviews = await db.interviews.find(
        {"candidate_id": {"$in": candidate_ids}},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    # Enrich with job and client info
    result = []
    for interview in interviews:
        job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
        client = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
        candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
        
        result.append({
            **interview,
            "job_title": job.get("title") if job else None,
            "job_description": job.get("description") if job else None,
            "job_location": job.get("location") if job else None,
            "job_work_model": job.get("work_model") if job else None,
            "company_name": client.get("company_name") if client else None,
            "candidate_name": candidate.get("name") if candidate else None
        })
    
    return result


@api_router.post("/candidate-portal/interviews/{interview_id}/book-slot")
async def candidate_portal_book_slot(
    interview_id: str,
    slot_id: str,
    current_candidate: dict = Depends(get_current_candidate)
):
    """Candidate books an interview slot (requires login)"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Verify this interview belongs to the candidate
    candidate_records = await db.candidates.find(
        {"$or": [
            {"candidate_portal_id": current_candidate["candidate_portal_id"]},
            {"email": current_candidate["email"]}
        ]},
        {"_id": 0, "candidate_id": 1}
    ).to_list(100)
    
    candidate_ids = [c["candidate_id"] for c in candidate_records]
    
    if interview["candidate_id"] not in candidate_ids:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="This interview does not belong to you"
        )
    
    # Verify interview is in correct status
    if interview["interview_status"] != "Awaiting Candidate Confirmation":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Interview is not awaiting confirmation (current status: {interview['interview_status']})"
        )
    
    # Find the selected slot
    selected_slot = None
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_id:
            selected_slot = slot
            break
    
    if not selected_slot:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Slot not found"
        )
    
    if not selected_slot.get("is_available", True):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Slot is no longer available"
        )
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update interview with selected slot
    update_data = {
        "selected_slot_id": slot_id,
        "scheduled_start_time": selected_slot["start_time"],
        "scheduled_end_time": selected_slot["end_time"],
        "interview_status": "Confirmed",
        "candidate_confirmation_timestamp": now,
        "updated_at": now
    }
    
    # Mark selected slot as unavailable
    updated_slots = []
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_id:
            slot["is_available"] = False
        updated_slots.append(slot)
    update_data["proposed_slots"] = updated_slots
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_data}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_candidate["candidate_portal_id"],
        user_email=current_candidate["email"],
        user_role="candidate",
        action_type="INTERVIEW_SLOT_BOOKED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        new_value={
            "slot_id": slot_id,
            "start_time": selected_slot["start_time"],
            "booked_via": "candidate_portal"
        }
    )
    
    # Create notification for recruiters
    notification_doc = {
        "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
        "type": "INTERVIEW_BOOKED",
        "title": f"Interview Confirmed: {current_candidate['name']}",
        "message": f"{current_candidate['name']} has confirmed their interview slot",
        "entity_type": "interview",
        "entity_id": interview_id,
        "client_id": interview["client_id"],
        "for_roles": ["admin", "recruiter"],
        "created_by": current_candidate["email"],
        "created_at": now,
        "read_by": []
    }
    await db.notifications.insert_one(notification_doc)
    
    # Send email notification to recruiters
    try:
        from notification_service import send_email, get_interview_booked_email_template
        
        job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
        client = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
        candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
        
        from datetime import datetime as dt
        slot_time = dt.fromisoformat(selected_slot["start_time"].replace('Z', '+00:00')).strftime('%B %d, %Y at %I:%M %p')
        
        subject, body = get_interview_booked_email_template(
            interview, candidate or {}, job or {}, client or {}, slot_time
        )
        
        recruiters = await db.users.find(
            {"role": {"$in": ["admin", "recruiter"]}},
            {"_id": 0, "email": 1}
        ).to_list(100)
        
        for recruiter in recruiters:
            await send_email(recruiter["email"], subject, body)
    except Exception as e:
        logging.error(f"Error sending interview booked notification: {str(e)}")
    
    return {"message": "Interview slot confirmed", "interview_id": interview_id}


# ============ CLIENT MANAGEMENT (Phase 2) ============

async def require_admin_or_recruiter(current_user: dict = Depends(get_current_user)):
    """Dependency to require admin or recruiter role"""
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Admin or recruiter access required"
        )
    return current_user

@api_router.get("/clients", response_model=list[ClientResponse])
async def list_clients(
    skip: int = 0,
    limit: int = 100,
    search: Optional[str] = None,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """List all clients with optional search and pagination"""
    query = {}
    if search:
        query["company_name"] = {"$regex": search, "$options": "i"}
    
    clients = await db.clients.find(query, {"_id": 0}).skip(skip).limit(limit).to_list(limit)
    
    # Get user count for each client
    result = []
    for client in clients:
        user_count = await db.users.count_documents({"client_id": client["client_id"]})
        result.append(ClientResponse(
            client_id=client["client_id"],
            company_name=client["company_name"],
            status=client["status"],
            created_at=client["created_at"],
            user_count=user_count
        ))
    
    return result

@api_router.post("/clients", response_model=ClientResponse)
async def create_client(
    client_data: ClientCreate,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Create a new client company"""
    # Generate unique client_id
    import uuid
    client_id = f"client_{uuid.uuid4().hex[:8]}"
    
    # Check if company name already exists
    existing = await db.clients.find_one({"company_name": client_data.company_name})
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Company name already exists"
        )
    
    client_doc = {
        "client_id": client_id,
        "company_name": client_data.company_name,
        "status": client_data.status,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": current_user["email"],
        "industry": client_data.industry,
        "website": client_data.website,
        "phone": client_data.phone,
        "address": client_data.address,
        "city": client_data.city,
        "state": client_data.state,
        "country": client_data.country,
        "postal_code": client_data.postal_code,
        "notes": client_data.notes
    }
    
    await db.clients.insert_one(client_doc)
    
    # Create default roles for this client
    await create_default_roles_for_client(client_id)
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="CLIENT_CREATE",
        entity_type="client",
        entity_id=client_id,
        new_value={"company_name": client_data.company_name, "status": client_data.status}
    )
    
    return ClientResponse(
        client_id=client_id,
        company_name=client_data.company_name,
        status=client_data.status,
        created_at=client_doc["created_at"],
        user_count=0,
        industry=client_data.industry,
        website=client_data.website,
        phone=client_data.phone,
        address=client_data.address,
        city=client_data.city,
        state=client_data.state,
        country=client_data.country,
        postal_code=client_data.postal_code,
        notes=client_data.notes
    )

@api_router.get("/clients/{client_id}", response_model=ClientResponse)
async def get_client(
    client_id: str,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Get a specific client by ID"""
    client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    user_count = await db.users.count_documents({"client_id": client_id})
    
    return ClientResponse(
        client_id=client["client_id"],
        company_name=client["company_name"],
        status=client["status"],
        created_at=client["created_at"],
        user_count=user_count,
        industry=client.get("industry"),
        website=client.get("website"),
        phone=client.get("phone"),
        address=client.get("address"),
        city=client.get("city"),
        state=client.get("state"),
        country=client.get("country"),
        postal_code=client.get("postal_code"),
        notes=client.get("notes")
    )

@api_router.put("/clients/{client_id}", response_model=ClientResponse)
async def update_client(
    client_id: str,
    client_data: ClientUpdate,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Update a client's information"""
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    update_data = client_data.model_dump(exclude_unset=True)
    if not update_data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No update data provided"
        )
    
    # Check company name uniqueness if updating
    if "company_name" in update_data:
        existing = await db.clients.find_one({
            "company_name": update_data["company_name"],
            "client_id": {"$ne": client_id}
        })
        if existing:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Company name already exists"
            )
    
    await db.clients.update_one(
        {"client_id": client_id},
        {"$set": update_data}
    )
    
    updated_client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
    user_count = await db.users.count_documents({"client_id": client_id})
    
    return ClientResponse(
        client_id=updated_client["client_id"],
        company_name=updated_client["company_name"],
        status=updated_client["status"],
        created_at=updated_client["created_at"],
        user_count=user_count,
        industry=updated_client.get("industry"),
        website=updated_client.get("website"),
        phone=updated_client.get("phone"),
        address=updated_client.get("address"),
        city=updated_client.get("city"),
        state=updated_client.get("state"),
        country=updated_client.get("country"),
        postal_code=updated_client.get("postal_code"),
        notes=updated_client.get("notes")
    )

@api_router.patch("/clients/{client_id}/disable")
async def disable_client(
    client_id: str,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Disable a client (soft delete)"""
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    await db.clients.update_one(
        {"client_id": client_id},
        {"$set": {"status": "inactive"}}
    )
    
    return {"message": "Client disabled successfully"}

@api_router.get("/clients/{client_id}/users", response_model=list[UserResponse])
async def list_client_users(
    client_id: str,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """List all users for a specific client"""
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    users = await db.users.find(
        {"client_id": client_id},
        {"_id": 0, "password_hash": 0}
    ).to_list(1000)
    
    return [UserResponse(
        email=user.get("email"),
        name=user.get("name", ""),
        role=user.get("role"),
        client_id=user.get("client_id"),
        phone=user.get("phone"),
        user_id=user.get("user_id")
    ) for user in users]

@api_router.post("/clients/{client_id}/users", response_model=UserResponse)
async def create_client_user(
    client_id: str,
    user_data: ClientUserCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Create a new user for a specific client. Sends welcome email with credentials."""
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    # Check if user already exists
    existing_user = await db.users.find_one({"email": user_data.email})
    if existing_user:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Email already registered"
        )
    
    # Hash password and create user
    password_hash = hash_password(user_data.password)
    user_id = f"user_{uuid.uuid4().hex[:12]}"
    
    user_doc = {
        "user_id": user_id,
        "email": user_data.email,
        "name": user_data.name,
        "phone": user_data.phone,
        "role": "client_user",  # Always client_user for this endpoint
        "client_id": client_id,
        "password_hash": password_hash,
        "must_change_password": True,  # Enforce password change on first login
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": current_user["email"]
    }
    
    await db.users.insert_one(user_doc)
    
    # Send welcome email with credentials
    frontend_url = os.environ.get('REACT_APP_FRONTEND_URL', 'https://arbeit.co.in')
    try:
        background_tasks.add_task(
            send_client_user_welcome_email,
            user_data.email,
            user_data.name,
            client.get("company_name", "Your Company"),
            user_data.password,  # The original password provided
            frontend_url
        )
    except Exception as e:
        print(f"Failed to send welcome email: {e}")
    
    return UserResponse(
        email=user_data.email,
        name=user_data.name,
        role="client_user",
        client_id=client_id,
        phone=user_data.phone,
        user_id=user_id
    )


@api_router.put("/clients/{client_id}/users/{user_email}", response_model=UserResponse)
async def update_client_user(
    client_id: str,
    user_email: str,
    user_data: ClientUserUpdate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Update a client user's information (name, phone, email). 
    If email is changed, sends account setup notification to new email."""
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    # Decode URL-encoded email
    from urllib.parse import unquote
    decoded_email = unquote(user_email)
    
    # Find the user
    user = await db.users.find_one({"email": decoded_email, "client_id": client_id})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found for this client"
        )
    
    # Build update data
    update_data = user_data.model_dump(exclude_unset=True)
    if not update_data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No update data provided"
        )
    
    # Check if email is being changed
    email_changed = False
    new_email = None
    temp_password = None
    
    if "email" in update_data and update_data["email"] != decoded_email:
        new_email = update_data["email"]
        # Check if new email already exists
        existing_user = await db.users.find_one({"email": new_email})
        if existing_user:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Email already in use by another account"
            )
        email_changed = True
        # Generate new temp password and set must_change_password flag
        temp_password = secrets.token_urlsafe(8)
        password_hash = hash_password(temp_password)
        update_data["password_hash"] = password_hash
        update_data["must_change_password"] = True
    
    await db.users.update_one(
        {"email": decoded_email, "client_id": client_id},
        {"$set": update_data}
    )
    
    # Fetch the updated user with the correct email
    final_email = new_email if email_changed else decoded_email
    updated_user = await db.users.find_one({"email": final_email}, {"_id": 0, "password_hash": 0})
    
    # Send notification if email changed
    if email_changed and new_email and temp_password:
        frontend_url = os.environ.get('REACT_APP_FRONTEND_URL', 'https://arbeit.co.in')
        try:
            background_tasks.add_task(
                send_client_user_welcome_email,
                new_email,
                updated_user.get("name", "User"),
                client.get("company_name", "Your Company"),
                temp_password,
                frontend_url
            )
        except Exception as e:
            print(f"Failed to send welcome email: {e}")
    
    return UserResponse(
        email=updated_user["email"],
        name=updated_user.get("name", ""),
        role=updated_user["role"],
        client_id=updated_user.get("client_id"),
        phone=updated_user.get("phone"),
        user_id=updated_user.get("user_id")
    )


@api_router.delete("/clients/{client_id}/users/{user_email}")
async def delete_client_user(
    client_id: str,
    user_email: str,
    current_user: dict = Depends(require_admin_or_recruiter)
):
    """Remove a user from a client"""
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    # Decode URL-encoded email
    from urllib.parse import unquote
    decoded_email = unquote(user_email)
    
    # Find the user
    user = await db.users.find_one({"email": decoded_email, "client_id": client_id})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found for this client"
        )
    
    # Don't allow deleting yourself
    if decoded_email == current_user["email"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot delete your own account"
        )
    
    # Delete the user
    await db.users.delete_one({"email": decoded_email, "client_id": client_id})
    
    return {"message": f"User {decoded_email} removed successfully"}


# ============ JOB REQUIREMENTS (Phase 3) ============

@api_router.post("/jobs", response_model=JobResponse)
async def create_job(
    job_data: JobCreate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Create a new job requirement"""
    import uuid
    
    # Check permission to create jobs
    if current_user["role"] == "client_user":
        has_permission = await check_permission(current_user, "can_create_jobs", current_user.get("client_id"))
        if not has_permission:
            # Log access denied
            await log_audit_event(
                user_id=current_user.get("user_id", current_user["email"]),
                user_email=current_user["email"],
                user_role=current_user["role"],
                action_type="ACCESS_DENIED",
                entity_type="job",
                client_id=current_user.get("client_id"),
                metadata={"required_permission": "can_create_jobs", "endpoint": "POST /jobs"}
            )
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_create_jobs required"
            )
    
    # Determine client_id based on user role
    if current_user["role"] == "client_user":
        # Client users can only create jobs for their own client
        client_id = current_user["client_id"]
    elif current_user["role"] in ["admin", "recruiter"]:
        # Admin/recruiter must specify client_id
        if not job_data.client_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="client_id is required for admin/recruiter"
            )
        client_id = job_data.client_id
    else:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Invalid role"
        )
    
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    # Generate job_id
    job_id = f"job_{uuid.uuid4().hex[:8]}"
    
    job_doc = {
        "job_id": job_id,
        "client_id": client_id,
        "title": job_data.title,
        "location": job_data.location,
        "employment_type": job_data.employment_type,
        "experience_range": job_data.experience_range.model_dump(),
        "salary_range": job_data.salary_range.model_dump() if job_data.salary_range else None,
        "work_model": job_data.work_model,
        "city": job_data.city,
        "notice_period_days": job_data.notice_period_days,
        "required_skills": job_data.required_skills,
        "description": job_data.description,
        "status": job_data.status,
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": current_user["email"]
    }
    
    await db.jobs.insert_one(job_doc)
    
    # Send notification to recruiters (in background)
    async def send_job_notifications():
        try:
            # Get all recruiters and admins to notify
            recruiters = await db.users.find(
                {"role": {"$in": ["admin", "recruiter"]}},
                {"_id": 0, "email": 1, "name": 1}
            ).to_list(100)
            
            # Generate email content
            subject, body = get_new_job_email_template(job_doc, client, current_user["email"])
            
            # Send email to each recruiter
            for recruiter in recruiters:
                result = await send_email(recruiter["email"], subject, body)
                if result["success"]:
                    logging.info(f"Job notification email sent to {recruiter['email']}")
                else:
                    logging.error(f"Failed to send job notification to {recruiter['email']}: {result.get('error')}")
            
            # Create in-app notification
            notification_doc = {
                "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
                "type": "NEW_JOB",
                "title": f"New Job: {job_data.title}",
                "message": f"{client['company_name']} has submitted a new job requirement for {job_data.title}",
                "entity_type": "job",
                "entity_id": job_id,
                "client_id": client_id,
                "for_roles": ["admin", "recruiter"],
                "created_by": current_user["email"],
                "created_at": datetime.now(timezone.utc).isoformat(),
                "read_by": []
            }
            await db.notifications.insert_one(notification_doc)
            
        except Exception as e:
            logging.error(f"Error sending job notifications: {str(e)}")
    
    # Run notifications in background
    background_tasks.add_task(send_job_notifications)
    
    return JobResponse(
        job_id=job_id,
        client_id=client_id,
        title=job_data.title,
        location=job_data.location,
        employment_type=job_data.employment_type,
        experience_range=job_data.experience_range,
        salary_range=job_data.salary_range,
        work_model=job_data.work_model,
        required_skills=job_data.required_skills,
        description=job_data.description,
        status=job_data.status,
        created_at=job_doc["created_at"],
        created_by=current_user["email"],
        company_name=client["company_name"]
    )

@api_router.get("/jobs", response_model=list[JobResponse])
async def list_jobs(
    skip: int = 0,
    limit: int = 100,
    search: Optional[str] = None,
    client_id: Optional[str] = None,
    status: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """List job requirements with tenant filtering"""
    # Check permission to view jobs
    if current_user["role"] == "client_user":
        has_permission = await check_permission(current_user, "can_view_jobs", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_jobs required"
            )
    
    query = {}
    
    # Tenant filtering
    if current_user["role"] == "client_user":
        # Client users can only see their own jobs
        query["client_id"] = current_user["client_id"]
    elif current_user["role"] in ["admin", "recruiter"]:
        # Admin/recruiter can filter by client_id or see all
        if client_id:
            query["client_id"] = client_id
    
    # Search by title or skills
    if search:
        query["$or"] = [
            {"title": {"$regex": search, "$options": "i"}},
            {"required_skills": {"$regex": search, "$options": "i"}}
        ]
    
    # Filter by status
    if status:
        query["status"] = status
    
    jobs = await db.jobs.find(query, {"_id": 0}).skip(skip).limit(limit).to_list(limit)
    
    # Populate company names
    result = []
    for job in jobs:
        client = await db.clients.find_one({"client_id": job["client_id"]})
        result.append(JobResponse(
            job_id=job["job_id"],
            client_id=job["client_id"],
            title=job["title"],
            location=job["location"],
            employment_type=job["employment_type"],
            experience_range=ExperienceRange(**job["experience_range"]),
            salary_range=SalaryRange(**job["salary_range"]) if job.get("salary_range") else None,
            work_model=job["work_model"],
            required_skills=job["required_skills"],
            description=job["description"],
            status=job["status"],
            created_at=job["created_at"],
            created_by=job["created_by"],
            company_name=client["company_name"] if client else None
        ))
    
    return result

@api_router.get("/jobs/{job_id}", response_model=JobResponse)
async def get_job(
    job_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get a specific job requirement"""
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to view jobs
        has_permission = await check_permission(current_user, "can_view_jobs", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_jobs required"
            )
    
    # Get client info
    client = await db.clients.find_one({"client_id": job["client_id"]})
    
    return JobResponse(
        job_id=job["job_id"],
        client_id=job["client_id"],
        title=job["title"],
        location=job["location"],
        employment_type=job["employment_type"],
        experience_range=ExperienceRange(**job["experience_range"]),
        salary_range=SalaryRange(**job["salary_range"]) if job.get("salary_range") else None,
        work_model=job["work_model"],
        required_skills=job["required_skills"],
        description=job["description"],
        status=job["status"],
        created_at=job["created_at"],
        created_by=job["created_by"],
        company_name=client["company_name"] if client else None
    )

@api_router.put("/jobs/{job_id}", response_model=JobResponse)
async def update_job(
    job_id: str,
    job_data: JobUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update a job requirement"""
    job = await db.jobs.find_one({"job_id": job_id})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to edit jobs
        has_permission = await check_permission(current_user, "can_edit_jobs", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_edit_jobs required"
            )
    
    update_data = job_data.model_dump(exclude_unset=True)
    if not update_data:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No update data provided"
        )
    
    # Convert nested models to dicts (if not already dicts)
    if "experience_range" in update_data:
        if hasattr(update_data["experience_range"], 'model_dump'):
            update_data["experience_range"] = update_data["experience_range"].model_dump()
    if "salary_range" in update_data and update_data["salary_range"]:
        if hasattr(update_data["salary_range"], 'model_dump'):
            update_data["salary_range"] = update_data["salary_range"].model_dump()
    
    await db.jobs.update_one(
        {"job_id": job_id},
        {"$set": update_data}
    )
    
    updated_job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    client = await db.clients.find_one({"client_id": updated_job["client_id"]})
    
    return JobResponse(
        job_id=updated_job["job_id"],
        client_id=updated_job["client_id"],
        title=updated_job["title"],
        location=updated_job["location"],
        employment_type=updated_job["employment_type"],
        experience_range=ExperienceRange(**updated_job["experience_range"]),
        salary_range=SalaryRange(**updated_job["salary_range"]) if updated_job.get("salary_range") else None,
        work_model=updated_job["work_model"],
        required_skills=updated_job["required_skills"],
        description=updated_job["description"],
        status=updated_job["status"],
        created_at=updated_job["created_at"],
        created_by=updated_job["created_by"],
        company_name=client["company_name"] if client else None
    )


@api_router.delete("/jobs/{job_id}")
async def delete_job(
    job_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a job (Admin only)"""
    # Only admin can delete
    if current_user["role"] not in ["admin"]:
        await log_audit_event(
            user_id=current_user.get("user_id", current_user["email"]),
            user_email=current_user["email"],
            user_role=current_user["role"],
            action_type="ACCESS_DENIED",
            entity_type="job",
            entity_id=job_id,
            metadata={"required_role": "admin", "attempted_action": "delete_job"}
        )
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only Arbeit Admin can delete jobs"
        )
    
    # Get job
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Delete all candidates for this job
    candidates = await db.candidates.find({"job_id": job_id}, {"_id": 0}).to_list(1000)
    candidate_count = len(candidates)
    
    for candidate in candidates:
        # Delete CV versions
        await db.candidate_cv_versions.delete_many({"candidate_id": candidate["candidate_id"]})
        # Delete reviews
        await db.candidate_reviews.delete_many({"candidate_id": candidate["candidate_id"]})
        # Delete candidate
        await db.candidates.delete_one({"candidate_id": candidate["candidate_id"]})
    
    # Delete the job
    await db.jobs.delete_one({"job_id": job_id})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="JOB_DELETE",
        entity_type="job",
        entity_id=job_id,
        client_id=job.get("client_id"),
        metadata={
            "job_title": job["title"],
            "candidates_deleted": candidate_count
        },
        previous_value={
            "title": job["title"],
            "client_id": job["client_id"],
            "status": job["status"]
        }
    )
    
    return {
        "message": "Job and all associated data deleted successfully",
        "job_id": job_id,
        "candidates_deleted": candidate_count
    }


@api_router.patch("/jobs/{job_id}/close")
async def close_job(
    job_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Close a job requirement"""
    job = await db.jobs.find_one({"job_id": job_id})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to edit jobs (closing is considered editing)
        has_permission = await check_permission(current_user, "can_edit_jobs", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_edit_jobs required"
            )
    
    await db.jobs.update_one(
        {"job_id": job_id},
        {"$set": {"status": "Closed"}}
    )
    
    return {"message": "Job closed successfully"}


# ============ CANDIDATE MANAGEMENT (Phase 4) ============

@api_router.post("/candidates/upload", response_model=CandidateResponse)
async def upload_candidate_cv(
    job_id: str = Form(...),
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload CV and create candidate with AI parsing"""
    # Check permission to upload CV
    if current_user["role"] == "client_user":
        has_permission = await check_permission(current_user, "can_upload_cv", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_upload_cv required"
            )
    elif current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin/recruiter can upload candidates"
        )
    
    # Verify job exists and get job details
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Generate candidate_id
    candidate_id = f"cand_{uuid.uuid4().hex[:8]}"
    
    # Save CV file
    cv_url = await save_cv_file(file, candidate_id)
    
    # Extract text from CV using proper PDF/DOCX parsing
    cv_text = await extract_text_from_cv(file)
    print(f"[DEBUG] Extracted CV text length: {len(cv_text)} chars")
    print(f"[DEBUG] CV text preview: {cv_text[:500]}")
    
    # Parse CV with AI
    parsed_resume = await parse_cv_with_ai(cv_text)
    
    # Generate candidate story with full parsed data
    candidate_data_for_story = {
        "name": parsed_resume.name,
        "current_role": parsed_resume.current_role,
        "skills": parsed_resume.skills,
        "experience": parsed_resume.experience,
        "education": parsed_resume.education,
        "summary": parsed_resume.summary
    }
    ai_story = await generate_candidate_story(candidate_data_for_story, job)
    
    # Create candidate document
    candidate_doc = {
        "candidate_id": candidate_id,
        "job_id": job_id,
        "name": parsed_resume.name,
        "current_role": parsed_resume.current_role,
        "email": parsed_resume.email,
        "phone": parsed_resume.phone,
        "linkedin": parsed_resume.linkedin,
        "skills": parsed_resume.skills,
        "experience": parsed_resume.experience,
        "education": parsed_resume.education,
        "summary": parsed_resume.summary,
        "cv_file_url": cv_url,
        "cv_text_original": cv_text,
        "cv_text_redacted": redact_text(cv_text),
        "ai_story": ai_story.model_dump(),
        "status": "NEW",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": current_user["email"]
    }
    
    await db.candidates.insert_one(candidate_doc)
    
    # Create initial CV version entry
    version_id = f"cv_v_{uuid.uuid4().hex[:12]}"
    version_doc = {
        "version_id": version_id,
        "candidate_id": candidate_id,
        "version_number": 1,
        "file_url": cv_url,
        "source_filename": file.filename,
        "uploaded_by_user_id": current_user.get("user_id", current_user["email"]),
        "uploaded_by_email": current_user["email"],
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "is_active": True,
        "ai_parsed_data": {
            "name": parsed_resume.name,
            "current_role": parsed_resume.current_role,
            "email": parsed_resume.email,
            "phone": parsed_resume.phone,
            "linkedin": parsed_resume.linkedin,
            "skills": parsed_resume.skills,
            "experience": parsed_resume.experience,
            "education": parsed_resume.education,
            "summary": parsed_resume.summary
        },
        "ai_story_json": ai_story.model_dump(),
        "fit_score": ai_story.fit_score,
        "deleted_at": None,
        "delete_type": None,
        "deleted_by_user_id": None
    }
    await db.candidate_cv_versions.insert_one(version_doc)
    
    return CandidateResponse(
        candidate_id=candidate_id,
        job_id=job_id,
        name=parsed_resume.name,
        current_role=parsed_resume.current_role,
        email=parsed_resume.email,
        phone=parsed_resume.phone,
        linkedin=parsed_resume.linkedin,
        skills=parsed_resume.skills,
        experience=parsed_resume.experience,
        education=parsed_resume.education,
        summary=parsed_resume.summary,
        cv_file_url=cv_url,
        ai_story=ai_story,
        status="NEW",
        created_at=candidate_doc["created_at"],
        created_by=current_user["email"]
    )

@api_router.post("/candidates", response_model=CandidateResponse)
async def create_candidate_manual(
    candidate_data: CandidateCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create candidate manually without CV upload"""
    # Check permission to create candidates
    if current_user["role"] == "client_user":
        has_permission = await check_permission(current_user, "can_create_candidates", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_create_candidates required"
            )
    elif current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin/recruiter can create candidates"
        )
    
    # Verify job exists
    job = await db.jobs.find_one({"job_id": candidate_data.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    candidate_id = f"cand_{uuid.uuid4().hex[:8]}"
    
    # Generate AI story
    ai_story = await generate_candidate_story(candidate_data.model_dump(), job)
    
    candidate_doc = {
        "candidate_id": candidate_id,
        "job_id": candidate_data.job_id,
        "name": candidate_data.name,
        "current_role": candidate_data.current_role,
        "email": candidate_data.email,
        "phone": candidate_data.phone,
        "skills": candidate_data.skills,
        "experience": candidate_data.experience,
        "education": candidate_data.education,
        "summary": candidate_data.summary,
        "cv_file_url": None,
        "ai_story": ai_story.model_dump(),
        "status": "NEW",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "created_by": current_user["email"]
    }
    
    await db.candidates.insert_one(candidate_doc)
    
    return CandidateResponse(
        candidate_id=candidate_id,
        job_id=candidate_data.job_id,
        name=candidate_data.name,
        current_role=candidate_data.current_role,
        email=candidate_data.email,
        phone=candidate_data.phone,
        linkedin=None,
        skills=candidate_data.skills,
        experience=candidate_data.experience,
        education=candidate_data.education,
        summary=candidate_data.summary,
        cv_file_url=None,
        ai_story=ai_story,
        status="NEW",
        created_at=candidate_doc["created_at"],
        created_by=current_user["email"]
    )

@api_router.get("/jobs/{job_id}/candidates", response_model=list[CandidateResponse])
async def list_job_candidates(
    job_id: str,
    show_rejected: bool = False,
    current_user: dict = Depends(get_current_user)
):
    """List all candidates for a job (excluding rejected by default)"""
    # Verify job exists and user has access
    job = await db.jobs.find_one({"job_id": job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to view candidates
        has_permission = await check_permission(current_user, "can_view_candidates", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_candidates required"
            )
    
    # Build query to exclude rejected candidates unless show_rejected is True
    query = {"job_id": job_id}
    if not show_rejected:
        query["status"] = {"$ne": "REJECT"}
    
    candidates = await db.candidates.find(
        query,
        {"_id": 0}
    ).to_list(1000)
    
    result = []
    for cand in candidates:
        result.append(CandidateResponse(
            candidate_id=cand["candidate_id"],
            job_id=cand["job_id"],
            name=cand["name"],
            current_role=cand.get("current_role"),
            email=cand.get("email"),
            phone=cand.get("phone"),
            linkedin=cand.get("linkedin"),
            skills=cand.get("skills", []),
            experience=cand.get("experience", []),
            education=cand.get("education", []),
            summary=cand.get("summary"),
            cv_file_url=cand.get("cv_file_url"),
            ai_story=CandidateStory(**cand["ai_story"]) if cand.get("ai_story") else None,
            status=cand["status"],
            created_at=cand["created_at"],
            created_by=cand["created_by"]
        ))
    
    return result

@api_router.get("/candidates/{candidate_id}", response_model=CandidateResponse)
async def get_candidate(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get candidate details"""
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to view candidates
        has_permission = await check_permission(current_user, "can_view_candidates", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_candidates required"
            )
    
    return CandidateResponse(
        candidate_id=candidate["candidate_id"],
        job_id=candidate["job_id"],
        name=candidate["name"],
        current_role=candidate.get("current_role"),
        email=candidate.get("email"),
        phone=candidate.get("phone"),
        linkedin=candidate.get("linkedin"),
        skills=candidate.get("skills", []),
        experience=candidate.get("experience", []),
        education=candidate.get("education", []),
        summary=candidate.get("summary"),
        cv_file_url=candidate.get("cv_file_url"),
        ai_story=CandidateStory(**candidate["ai_story"]) if candidate.get("ai_story") else None,
        status=candidate["status"],
        created_at=candidate["created_at"],
        created_by=candidate["created_by"]
    )


@api_router.delete("/candidates/{candidate_id}")
async def delete_candidate(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a candidate (Admin only)"""
    # Only admin can delete
    if current_user["role"] not in ["admin"]:
        await log_audit_event(
            user_id=current_user.get("user_id", current_user["email"]),
            user_email=current_user["email"],
            user_role=current_user["role"],
            action_type="ACCESS_DENIED",
            entity_type="candidate",
            entity_id=candidate_id,
            metadata={"required_role": "admin", "attempted_action": "delete_candidate"}
        )
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only Arbeit Admin can delete candidates"
        )
    
    # Get candidate
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Get job for client_id
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    
    # Delete CV versions
    cv_versions = await db.candidate_cv_versions.find({"candidate_id": candidate_id}).to_list(1000)
    await db.candidate_cv_versions.delete_many({"candidate_id": candidate_id})
    
    # Delete reviews
    reviews = await db.candidate_reviews.find({"candidate_id": candidate_id}).to_list(1000)
    await db.candidate_reviews.delete_many({"candidate_id": candidate_id})
    
    # Delete candidate
    await db.candidates.delete_one({"candidate_id": candidate_id})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="CANDIDATE_DELETE",
        entity_type="candidate",
        entity_id=candidate_id,
        client_id=job.get("client_id") if job else None,
        metadata={
            "candidate_name": candidate["name"],
            "job_id": candidate["job_id"],
            "cv_versions_deleted": len(cv_versions),
            "reviews_deleted": len(reviews)
        },
        previous_value={
            "name": candidate["name"],
            "job_id": candidate["job_id"],
            "status": candidate["status"]
        }
    )
    
    return {
        "message": "Candidate and all associated data deleted successfully",
        "candidate_id": candidate_id,
        "cv_versions_deleted": len(cv_versions),
        "reviews_deleted": len(reviews)
    }


@api_router.get("/candidates/{candidate_id}/cv")
async def get_candidate_cv(
    candidate_id: str,
    redacted: bool = True,
    current_user: dict = Depends(get_current_user)
):
    """Get candidate CV (redacted or full)"""
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permission to view CV
        if redacted:
            has_permission = await check_permission(current_user, "can_view_redacted_cv", current_user.get("client_id"))
        else:
            has_permission = await check_permission(current_user, "can_view_full_cv", current_user.get("client_id"))
        
        if not has_permission:
            permission_type = "can_view_redacted_cv" if redacted else "can_view_full_cv"
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"Permission denied: {permission_type} required"
            )
        # Client users always get redacted version unless they have full CV permission
        if not await check_permission(current_user, "can_view_full_cv", current_user.get("client_id")):
            redacted = True
    
    cv_text = candidate.get("cv_text_redacted") if redacted else candidate.get("cv_text_original")
    
    return {
        "candidate_id": candidate_id,
        "cv_text": cv_text,
        "is_redacted": redacted
    }

@api_router.put("/candidates/{candidate_id}", response_model=CandidateResponse)
async def update_candidate(
    candidate_id: str,
    update_data: CandidateUpdate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Update candidate information"""
    candidate = await db.candidates.find_one({"candidate_id": candidate_id})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Store old status for comparison
    old_status = candidate.get("status")
    
    # Verify access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        # Check permissions for client users
        update_dict = update_data.model_dump(exclude_unset=True)
        
        # Check if updating status
        if "status" in update_dict:
            has_permission = await check_permission(current_user, "can_update_candidate_status", current_user.get("client_id"))
            if not has_permission:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Permission denied: can_update_candidate_status required"
                )
        
        # Check if updating other fields
        other_fields = set(update_dict.keys()) - {"status"}
        if other_fields:
            has_permission = await check_permission(current_user, "can_edit_candidates", current_user.get("client_id"))
            if not has_permission:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Permission denied: can_edit_candidates required"
                )
    
    update_dict = update_data.model_dump(exclude_unset=True)
    if not update_dict:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No update data provided"
        )
    
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {"$set": update_dict}
    )
    
    updated_candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    
    # Trigger notification if status changed
    new_status = updated_candidate.get("status")
    if "status" in update_dict and old_status != new_status:
        # Send notification in background
        background_tasks.add_task(
            send_candidate_status_change_notification,
            candidate_id=candidate_id,
            old_status=old_status,
            new_status=new_status,
            changed_by=current_user["email"]
        )
    
    return CandidateResponse(
        candidate_id=updated_candidate["candidate_id"],
        job_id=updated_candidate["job_id"],
        name=updated_candidate["name"],
        current_role=updated_candidate.get("current_role"),
        email=updated_candidate.get("email"),
        phone=updated_candidate.get("phone"),
        linkedin=updated_candidate.get("linkedin"),
        skills=updated_candidate.get("skills", []),
        experience=updated_candidate.get("experience", []),
        education=updated_candidate.get("education", []),
        summary=updated_candidate.get("summary"),
        cv_file_url=updated_candidate.get("cv_file_url"),
        ai_story=CandidateStory(**updated_candidate["ai_story"]) if updated_candidate.get("ai_story") else None,
        status=updated_candidate["status"],
        created_at=updated_candidate["created_at"],
        created_by=updated_candidate["created_by"]
    )

@api_router.post("/candidates/{candidate_id}/regenerate-story", response_model=CandidateResponse)
async def regenerate_candidate_story(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Regenerate AI candidate story"""
    # Only admin/recruiter can regenerate
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin/recruiter can regenerate stories"
        )
    
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    
    # Generate new story
    ai_story = await generate_candidate_story(candidate, job)
    
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {"$set": {"ai_story": ai_story.model_dump()}}
    )
    
    updated_candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    
    return CandidateResponse(
        candidate_id=updated_candidate["candidate_id"],
        job_id=updated_candidate["job_id"],
        name=updated_candidate["name"],
        current_role=updated_candidate.get("current_role"),
        email=updated_candidate.get("email"),
        phone=updated_candidate.get("phone"),
        linkedin=updated_candidate.get("linkedin"),
        skills=updated_candidate.get("skills", []),
        experience=updated_candidate.get("experience", []),
        education=updated_candidate.get("education", []),
        summary=updated_candidate.get("summary"),
        cv_file_url=updated_candidate.get("cv_file_url"),
        ai_story=ai_story,
        status=updated_candidate["status"],
        created_at=updated_candidate["created_at"],
        created_by=updated_candidate["created_by"]
    )


# ============ REVIEW WORKFLOW (Phase 5) ============

@api_router.post("/candidates/{candidate_id}/review", response_model=ReviewResponse)
async def create_review(
    candidate_id: str,
    review_data: ReviewCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a review entry for a candidate"""
    # Verify candidate exists and user has access
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access (tenant check for client users)
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Generate review ID
    review_id = f"rev_{uuid.uuid4().hex[:8]}"
    
    # Create review document
    review_doc = {
        "review_id": review_id,
        "candidate_id": candidate_id,
        "user_id": current_user["email"],
        "user_name": current_user["name"],
        "user_role": current_user["role"],
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "action": review_data.action,
        "comment": review_data.comment
    }
    
    await db.reviews.insert_one(review_doc)
    
    # Update candidate status if action is not COMMENT
    if review_data.action != "COMMENT":
        new_status = review_data.action  # APPROVE, PIPELINE, or REJECT
        await db.candidates.update_one(
            {"candidate_id": candidate_id},
            {"$set": {"status": new_status}}
        )
    
    return ReviewResponse(
        review_id=review_id,
        candidate_id=candidate_id,
        user_id=current_user["email"],
        user_name=current_user["name"],
        user_role=current_user["role"],
        timestamp=review_doc["timestamp"],
        action=review_data.action,
        comment=review_data.comment
    )

@api_router.get("/candidates/{candidate_id}/reviews", response_model=list[ReviewResponse])
async def list_candidate_reviews(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all reviews for a candidate (newest first)"""
    # Verify candidate exists and user has access
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Get reviews sorted by timestamp descending (newest first)
    reviews = await db.reviews.find(
        {"candidate_id": candidate_id},
        {"_id": 0}
    ).sort("timestamp", -1).to_list(1000)
    
    return [
        ReviewResponse(
            review_id=review["review_id"],
            candidate_id=review["candidate_id"],
            user_id=review["user_id"],
            user_name=review["user_name"],
            user_role=review["user_role"],
            timestamp=review["timestamp"],
            action=review["action"],
            comment=review.get("comment")
        )
        for review in reviews
    ]


# ============ STORY VIEW & PDF EXPORT (Phase 6) ============

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfgen import canvas
from io import BytesIO

@api_router.post("/candidates/{candidate_id}/story/regenerate")
async def regenerate_candidate_story_endpoint(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Regenerate AI candidate story with editorial formatting"""
    # Only admin and recruiter can regenerate
    if current_user["role"] not in ["admin", "recruiter"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only admin and recruiter can regenerate stories"
        )
    
    # Get candidate
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Generate new story
    new_story = await generate_candidate_story(candidate, job)
    
    # Update candidate with new story and timestamp
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {
            "$set": {
                "ai_story": new_story.model_dump(),
                "story_last_generated": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    # Return updated candidate
    updated_candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    ai_story = CandidateStory(**updated_candidate["ai_story"]) if updated_candidate.get("ai_story") else None
    
    return CandidateResponse(
        candidate_id=updated_candidate["candidate_id"],
        job_id=updated_candidate["job_id"],
        name=updated_candidate["name"],
        current_role=updated_candidate.get("current_role"),
        email=updated_candidate.get("email"),
        phone=updated_candidate.get("phone"),
        linkedin=updated_candidate.get("linkedin"),
        skills=updated_candidate.get("skills", []),
        experience=updated_candidate.get("experience", []),
        education=updated_candidate.get("education", []),
        summary=updated_candidate.get("summary"),
        cv_file_url=updated_candidate.get("cv_file_url"),
        ai_story=ai_story,
        status=updated_candidate["status"],
        created_at=updated_candidate["created_at"],
        created_by=updated_candidate["created_by"]
    )

@api_router.get("/candidates/{candidate_id}/story/export")
async def export_candidate_story_pdf(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Export candidate story as PDF"""
    # Get candidate
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Generate PDF
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=25, bottomMargin=25, leftMargin=25, rightMargin=25)
    
    # Container for PDF elements
    story_elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=28,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=8,
        alignment=TA_CENTER,
        fontName='Times-Bold'
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontSize=16,
        textColor=colors.HexColor('#7F8C8D'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica'
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#D4AF37'),
        spaceAfter=12,
        spaceBefore=16,
        fontName='Times-Bold',
        borderWidth=1,
        borderColor=colors.HexColor('#D4AF37'),
        borderPadding=8
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=8,
        fontName='Helvetica',
        leading=16
    )
    
    # Hero section
    story_elements.append(Paragraph(candidate["name"], title_style))
    if candidate.get("current_role"):
        story_elements.append(Paragraph(candidate["current_role"], subtitle_style))
    
    # Status and fit score
    ai_story = candidate.get("ai_story", {})
    fit_score = ai_story.get("fit_score", 0) if ai_story else 0
    status_text = f"Status: {candidate['status']} | Fit Score: {fit_score}%"
    story_elements.append(Paragraph(status_text, body_style))
    story_elements.append(Spacer(1, 0.3*inch))
    
    # Professional Summary
    if ai_story and ai_story.get("summary"):
        story_elements.append(Paragraph("Professional Summary", heading_style))
        story_elements.append(Paragraph(ai_story["summary"], body_style))
        story_elements.append(Spacer(1, 0.2*inch))
    
    # Skills
    if candidate.get("skills"):
        story_elements.append(Paragraph("Core Skills", heading_style))
        skills_text = " • ".join(candidate["skills"][:10])
        story_elements.append(Paragraph(skills_text, body_style))
        story_elements.append(Spacer(1, 0.2*inch))
    
    # Experience
    if candidate.get("experience"):
        story_elements.append(Paragraph("Career Timeline", heading_style))
        for exp in candidate["experience"][:5]:
            role_title = exp.get("role", "Position")
            company = exp.get("company", "")
            duration = exp.get("duration", "")
            exp_text = f"<b>{role_title}</b>"
            if company:
                exp_text += f" at {company}"
            if duration:
                exp_text += f" ({duration})"
            story_elements.append(Paragraph(exp_text, body_style))
            
            if exp.get("achievements") and isinstance(exp["achievements"], list):
                for achievement in exp["achievements"][:3]:
                    story_elements.append(Paragraph(f"  • {achievement}", body_style))
            story_elements.append(Spacer(1, 0.1*inch))
    
    # Highlights
    if ai_story and ai_story.get("highlights"):
        story_elements.append(Paragraph("Key Achievements", heading_style))
        for highlight in ai_story["highlights"][:5]:
            story_elements.append(Paragraph(f"✓ {highlight}", body_style))
        story_elements.append(Spacer(1, 0.2*inch))
    
    # Education
    if candidate.get("education"):
        story_elements.append(Paragraph("Education", heading_style))
        for edu in candidate["education"][:3]:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            year = edu.get("year", "")
            edu_text = f"{degree}"
            if institution:
                edu_text += f" - {institution}"
            if year:
                edu_text += f" ({year})"
            story_elements.append(Paragraph(edu_text, body_style))
        story_elements.append(Spacer(1, 0.3*inch))
    
    # Footer
    footer_text = f"Generated by Arbeit Talent Platform – {datetime.now().strftime('%B %d, %Y')} – Candidate ID: {candidate_id}"
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.HexColor('#95A5A6'),
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique'
    )
    story_elements.append(Spacer(1, 0.5*inch))
    story_elements.append(Paragraph(footer_text, footer_style))
    
    # Build PDF
    doc.build(story_elements)
    buffer.seek(0)
    
    # Return PDF as response
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"attachment; filename=candidate_story_{candidate['name'].replace(' ', '_')}_{candidate_id}.pdf"
        }
    )


# ============ HEALTH CHECK ============

@api_router.get("/")
async def root():
    return {"message": "Arbeit Talent Portal API", "version": "1.0.0"}





# ============ CV VERSIONING & REPLACEMENT ============

@api_router.post("/candidates/{candidate_id}/cv", response_model=CVVersionResponse)
async def replace_candidate_cv(
    candidate_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload a new CV for an existing candidate (replacement workflow)"""
    
    # Get candidate
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access and tenant isolation
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        # Check permission to replace CV
        has_permission = await check_permission(current_user, "can_replace_cv", current_user.get("client_id"))
        if not has_permission:
            # Log access denied
            await log_audit_event(
                user_id=current_user.get("user_id", current_user["email"]),
                user_email=current_user["email"],
                user_role=current_user["role"],
                action_type="ACCESS_DENIED",
                entity_type="candidate_cv",
                entity_id=candidate_id,
                client_id=current_user.get("client_id"),
                metadata={"required_permission": "can_replace_cv", "endpoint": "POST /candidates/{candidate_id}/cv"}
            )
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_replace_cv required"
            )
    
    # Get current version number (find max version_number for this candidate)
    existing_versions = await db.candidate_cv_versions.find(
        {"candidate_id": candidate_id}
    ).to_list(1000)
    
    next_version_number = max([v.get("version_number", 0) for v in existing_versions], default=0) + 1
    
    # Mark current active version as inactive
    if existing_versions:
        await db.candidate_cv_versions.update_many(
            {"candidate_id": candidate_id, "is_active": True},
            {"$set": {"is_active": False}}
        )
    
    # Save new CV file
    version_id = f"cv_v_{uuid.uuid4().hex[:12]}"
    file_extension = Path(file.filename).suffix
    filename = f"{candidate_id}_v{next_version_number}{file_extension}"
    file_path = UPLOAD_DIR / filename
    
    with open(file_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    cv_url = f"/api/uploads/{filename}"
    
    # Extract text from CV using proper PDF/DOCX parsing
    cv_text = await extract_text_from_cv(file)
    print(f"[DEBUG] Replace CV - Extracted text length: {len(cv_text)} chars")
    
    # Parse CV with AI
    parsed_resume = await parse_cv_with_ai(cv_text)
    
    # Generate candidate story with full parsed data
    candidate_data_for_story = {
        "name": parsed_resume.name or candidate.get("name"),
        "current_role": parsed_resume.current_role or candidate.get("current_role"),
        "skills": parsed_resume.skills or candidate.get("skills", []),
        "experience": parsed_resume.experience or candidate.get("experience", []),
        "education": parsed_resume.education or candidate.get("education", []),
        "summary": parsed_resume.summary or candidate.get("summary", "")
    }
    ai_story = await generate_candidate_story(candidate_data_for_story, job)
    
    # Create new version entry
    version_doc = {
        "version_id": version_id,
        "candidate_id": candidate_id,
        "version_number": next_version_number,
        "file_url": cv_url,
        "source_filename": file.filename,
        "uploaded_by_user_id": current_user.get("user_id", current_user["email"]),
        "uploaded_by_email": current_user["email"],
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
        "is_active": True,
        "ai_parsed_data": {
            "name": parsed_resume.name,
            "current_role": parsed_resume.current_role,
            "email": parsed_resume.email,
            "phone": parsed_resume.phone,
            "linkedin": parsed_resume.linkedin,
            "skills": parsed_resume.skills,
            "experience": parsed_resume.experience,
            "education": parsed_resume.education,
            "summary": parsed_resume.summary
        },
        "ai_story_json": ai_story.model_dump(),
        "fit_score": ai_story.fit_score,
        "deleted_at": None,
        "delete_type": None,
        "deleted_by_user_id": None
    }
    
    await db.candidate_cv_versions.insert_one(version_doc)
    
    # Update main candidate document with new data
    await db.candidates.update_one(
        {"candidate_id": candidate_id},
        {
            "$set": {
                "name": parsed_resume.name or candidate.get("name"),
                "current_role": parsed_resume.current_role or candidate.get("current_role"),
                "email": parsed_resume.email or candidate.get("email"),
                "phone": parsed_resume.phone or candidate.get("phone"),
                "linkedin": parsed_resume.linkedin or candidate.get("linkedin"),
                "skills": parsed_resume.skills or candidate.get("skills", []),
                "experience": parsed_resume.experience or candidate.get("experience", []),
                "education": parsed_resume.education or candidate.get("education", []),
                "summary": parsed_resume.summary or candidate.get("summary"),
                "cv_file_url": cv_url,
                "cv_text_original": cv_text,
                "cv_text_redacted": redact_text(cv_text),
                "ai_story": ai_story.model_dump(),
                "story_last_generated": datetime.now(timezone.utc).isoformat()
            }
        }
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="CV_REPLACED",
        entity_type="candidate_cv",
        entity_id=candidate_id,
        client_id=job.get("client_id"),
        metadata={
            "version_id": version_id,
            "version_number": next_version_number,
            "filename": file.filename,
            "previous_versions": len(existing_versions)
        },
        previous_value={"cv_file_url": candidate.get("cv_file_url")},
        new_value={"cv_file_url": cv_url, "version_number": next_version_number}
    )
    
    return CVVersionResponse(**version_doc)

@api_router.get("/candidates/{candidate_id}/cv/versions", response_model=list[CVVersionListItem])
async def list_cv_versions(
    candidate_id: str,
    include_deleted: bool = False,
    current_user: dict = Depends(get_current_user)
):
    """List all CV versions for a candidate"""
    
    # Get candidate to check access
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify job access
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        # Check permission to view candidates
        has_permission = await check_permission(current_user, "can_view_candidates", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_candidates required"
            )
    
    # Build query
    query = {"candidate_id": candidate_id}
    
    # Filter out deleted versions for non-admin users
    if not include_deleted or current_user["role"] not in ["admin", "recruiter"]:
        query["deleted_at"] = None
    
    versions = await db.candidate_cv_versions.find(
        query,
        {"_id": 0}
    ).sort("version_number", -1).to_list(1000)
    
    return [CVVersionListItem(
        version_id=v["version_id"],
        version_number=v["version_number"],
        source_filename=v["source_filename"],
        uploaded_by_email=v["uploaded_by_email"],
        uploaded_at=v["uploaded_at"],
        is_active=v["is_active"],
        deleted_at=v.get("deleted_at"),
        delete_type=v.get("delete_type")
    ) for v in versions]

@api_router.get("/candidates/{candidate_id}/cv/versions/{version_id}")
async def get_cv_version_file(
    candidate_id: str,
    version_id: str,
    redacted: bool = True,
    current_user: dict = Depends(get_current_user)
):
    """View a specific CV version (file content)"""
    
    # Get version
    version = await db.candidate_cv_versions.find_one(
        {"version_id": version_id, "candidate_id": candidate_id},
        {"_id": 0}
    )
    if not version:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV version not found"
        )
    
    # Get candidate to check access
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        
        # Check CV viewing permissions
        if redacted:
            has_permission = await check_permission(current_user, "can_view_redacted_cv", current_user.get("client_id"))
        else:
            has_permission = await check_permission(current_user, "can_view_full_cv", current_user.get("client_id"))
        
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"Permission denied: {'can_view_redacted_cv' if redacted else 'can_view_full_cv'} required"
            )
        
        # Force redacted for client users without full CV permission
        if not await check_permission(current_user, "can_view_full_cv", current_user.get("client_id")):
            redacted = True
    
    # Return version details (could extend to serve file content if needed)
    return {
        "version_id": version["version_id"],
        "version_number": version["version_number"],
        "file_url": version["file_url"],
        "source_filename": version["source_filename"],
        "uploaded_at": version["uploaded_at"],
        "uploaded_by": version["uploaded_by_email"],
        "is_active": version["is_active"],
        "ai_parsed_data": version.get("ai_parsed_data") if not redacted else None,
        "fit_score": version.get("fit_score")
    }

@api_router.delete("/candidates/{candidate_id}/cv/versions/{version_id}")
async def delete_cv_version(
    candidate_id: str,
    version_id: str,
    mode: str = "soft",  # "soft" or "hard"
    current_user: dict = Depends(get_current_user)
):
    """Delete a CV version (Admin only)"""
    
    # Only admin can delete versions
    if current_user["role"] not in ["admin"]:
        await log_audit_event(
            user_id=current_user.get("user_id", current_user["email"]),
            user_email=current_user["email"],
            user_role=current_user["role"],
            action_type="ACCESS_DENIED",
            entity_type="candidate_cv",
            entity_id=version_id,
            metadata={"required_role": "admin", "attempted_action": "delete_cv_version"}
        )
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only Arbeit Admin can delete CV versions"
        )
    
    # Get version
    version = await db.candidate_cv_versions.find_one(
        {"version_id": version_id, "candidate_id": candidate_id},
        {"_id": 0}
    )
    if not version:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="CV version not found"
        )
    
    # Cannot delete active version
    if version["is_active"]:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot delete active CV version. Upload a new version first."
        )
    
    if mode == "soft":
        # Soft delete: mark as deleted but keep data
        await db.candidate_cv_versions.update_one(
            {"version_id": version_id},
            {
                "$set": {
                    "deleted_at": datetime.now(timezone.utc).isoformat(),
                    "delete_type": "soft",
                    "deleted_by_user_id": current_user["email"]
                }
            }
        )
        
        action_type = "CV_SOFT_DELETE"
        message = "CV version soft deleted (archived)"
        
    elif mode == "hard":
        # Hard delete: remove file and mark as hard deleted
        file_path = UPLOAD_DIR / Path(version["file_url"]).name
        if file_path.exists():
            file_path.unlink()
        
        await db.candidate_cv_versions.update_one(
            {"version_id": version_id},
            {
                "$set": {
                    "deleted_at": datetime.now(timezone.utc).isoformat(),
                    "delete_type": "hard",
                    "deleted_by_user_id": current_user["email"],
                    "file_url": "[HARD_DELETED]"
                }
            }
        )
        
        action_type = "CV_HARD_DELETE"
        message = "CV version hard deleted (file removed)"
    else:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid mode. Use 'soft' or 'hard'"
        )
    
    # Get candidate for client_id
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type=action_type,
        entity_type="candidate_cv",
        entity_id=version_id,
        client_id=job.get("client_id"),
        metadata={
            "candidate_id": candidate_id,
            "version_number": version["version_number"],
            "delete_mode": mode,
            "filename": version["source_filename"]
        },
        previous_value={
            "file_url": version["file_url"],
            "version_number": version["version_number"]
        }
    )
    
    return {
        "message": message,
        "version_id": version_id,
        "version_number": version["version_number"],
        "delete_type": mode
    }


# ============ GOVERNANCE ROUTES (RBAC) ============

@api_router.get("/governance/roles", response_model=list[ClientRoleResponse])
async def list_client_roles(
    client_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """List all roles for a client (admin sees all clients, client_user sees their own)"""
    # Admin can see all clients or filter by client_id
    if current_user["role"] in ["admin", "recruiter"]:
        query = {"client_id": client_id} if client_id else {}
    else:
        # Client user can only see their own client's roles
        query = {"client_id": current_user["client_id"]}
    
    roles = await db.client_roles.find(query, {"_id": 0}).to_list(1000)
    return [ClientRoleResponse(**role) for role in roles]

@api_router.post("/governance/roles", response_model=ClientRoleResponse)
async def create_client_role(
    client_id: str,
    role_data: ClientRoleCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a new role for a client"""
    # Check permission
    if current_user["role"] not in ["admin", "recruiter"]:
        has_permission = await check_permission(current_user, "can_manage_roles", client_id)
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_manage_roles required"
            )
    
    # Verify client exists
    client = await db.clients.find_one({"client_id": client_id})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    role_id = f"role_{uuid.uuid4().hex[:12]}"
    role_doc = {
        "role_id": role_id,
        "client_id": client_id,
        "name": role_data.name,
        "description": role_data.description,
        "permissions": role_data.permissions.model_dump(),
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.client_roles.insert_one(role_doc)
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ROLE_CREATE",
        entity_type="role",
        entity_id=role_id,
        client_id=client_id,
        new_value={"name": role_data.name, "permissions": role_data.permissions.model_dump()}
    )
    
    return ClientRoleResponse(**role_doc)

@api_router.put("/governance/roles/{role_id}", response_model=ClientRoleResponse)
async def update_client_role(
    role_id: str,
    role_data: ClientRoleUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update a role's name, description, or permissions"""
    role = await db.client_roles.find_one({"role_id": role_id}, {"_id": 0})
    if not role:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Role not found"
        )
    
    # Check permission
    if current_user["role"] not in ["admin", "recruiter"]:
        has_permission = await check_permission(current_user, "can_manage_roles", role["client_id"])
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_manage_roles required"
            )
    
    # Prepare update
    update_data = role_data.model_dump(exclude_unset=True)
    if "permissions" in update_data and update_data["permissions"]:
        update_data["permissions"] = update_data["permissions"].model_dump()
    
    update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    # Update role
    await db.client_roles.update_one(
        {"role_id": role_id},
        {"$set": update_data}
    )
    
    updated_role = await db.client_roles.find_one({"role_id": role_id}, {"_id": 0})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ROLE_UPDATE",
        entity_type="role",
        entity_id=role_id,
        client_id=role["client_id"],
        previous_value={"permissions": role.get("permissions")},
        new_value={"permissions": updated_role.get("permissions")}
    )
    
    return ClientRoleResponse(**updated_role)

@api_router.delete("/governance/roles/{role_id}")
async def delete_client_role(
    role_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a role (admin only or client owner with can_manage_roles)"""
    role = await db.client_roles.find_one({"role_id": role_id}, {"_id": 0})
    if not role:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Role not found"
        )
    
    # Check permission
    if current_user["role"] not in ["admin", "recruiter"]:
        has_permission = await check_permission(current_user, "can_manage_roles", role["client_id"])
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_manage_roles required"
            )
    
    # Remove all user assignments for this role
    await db.user_client_roles.delete_many({"client_role_id": role_id})
    
    # Delete the role
    await db.client_roles.delete_one({"role_id": role_id})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ROLE_DELETE",
        entity_type="role",
        entity_id=role_id,
        client_id=role["client_id"],
        previous_value={"name": role["name"], "permissions": role.get("permissions")}
    )
    
    return {"message": "Role deleted successfully"}

# User Role Assignment Endpoints

@api_router.post("/governance/user-roles", response_model=UserRoleResponse)
async def assign_role_to_user(
    assignment: UserRoleAssignment,
    current_user: dict = Depends(get_current_user)
):
    """Assign a role to a user"""
    # Get the role to find its client_id
    role = await db.client_roles.find_one({"role_id": assignment.client_role_id}, {"_id": 0})
    if not role:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Role not found"
        )
    
    # Check permission
    if current_user["role"] not in ["admin", "recruiter"]:
        has_permission = await check_permission(current_user, "can_manage_users", role["client_id"])
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_manage_users required"
            )
    
    # Get user
    user = await db.users.find_one({"email": assignment.user_id}, {"_id": 0})
    if not user:
        user = await db.users.find_one({"user_id": assignment.user_id}, {"_id": 0})
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    # Check if assignment already exists
    existing = await db.user_client_roles.find_one({
        "user_id": assignment.user_id,
        "client_role_id": assignment.client_role_id
    })
    if existing:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Role already assigned to user"
        )
    
    assignment_id = f"assignment_{uuid.uuid4().hex[:12]}"
    assignment_doc = {
        "assignment_id": assignment_id,
        "user_id": assignment.user_id,
        "user_email": user["email"],
        "client_id": role["client_id"],
        "client_role_id": assignment.client_role_id,
        "role_name": role["name"],
        "assigned_by": current_user["email"],
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    await db.user_client_roles.insert_one(assignment_doc)
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ROLE_ASSIGN",
        entity_type="user_role",
        entity_id=assignment_id,
        client_id=role["client_id"],
        new_value={"user_id": assignment.user_id, "role_name": role["name"]}
    )
    
    return UserRoleResponse(**assignment_doc)

@api_router.get("/governance/user-roles", response_model=list[UserRoleResponse])
async def list_user_role_assignments(
    client_id: Optional[str] = None,
    user_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """List user role assignments"""
    # Build query
    query = {}
    
    if current_user["role"] in ["admin", "recruiter"]:
        if client_id:
            query["client_id"] = client_id
        if user_id:
            query["user_id"] = user_id
    else:
        # Client user can only see their own client
        query["client_id"] = current_user["client_id"]
    
    assignments = await db.user_client_roles.find(query, {"_id": 0}).to_list(1000)
    return [UserRoleResponse(**a) for a in assignments]

@api_router.delete("/governance/user-roles/{assignment_id}")
async def revoke_role_from_user(
    assignment_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Revoke a role assignment from a user"""
    assignment = await db.user_client_roles.find_one({"assignment_id": assignment_id}, {"_id": 0})
    if not assignment:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Assignment not found"
        )
    
    # Check permission
    if current_user["role"] not in ["admin", "recruiter"]:
        has_permission = await check_permission(current_user, "can_manage_users", assignment["client_id"])
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_manage_users required"
            )
    
    await db.user_client_roles.delete_one({"assignment_id": assignment_id})
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ROLE_REVOKE",
        entity_type="user_role",
        entity_id=assignment_id,
        client_id=assignment["client_id"],
        previous_value={"user_id": assignment["user_id"], "role_name": assignment["role_name"]}
    )
    
    return {"message": "Role revoked successfully"}

# Audit Log Endpoints

@api_router.get("/governance/audit", response_model=list[AuditLogEntry])
async def get_audit_logs(
    client_id: Optional[str] = None,
    user_id: Optional[str] = None,
    action_type: Optional[str] = None,
    entity_type: Optional[str] = None,
    from_date: Optional[str] = None,
    to_date: Optional[str] = None,
    limit: int = 100,
    skip: int = 0,
    current_user: dict = Depends(get_current_user)
):
    """Get audit logs with filtering"""
    # Build query
    query = {}
    
    # Permission check
    if current_user["role"] in ["admin", "recruiter"]:
        # Admin can see all logs, optionally filtered by client
        if client_id:
            query["client_id"] = client_id
    else:
        # Client user needs can_view_audit_log permission
        has_permission = await check_permission(current_user, "can_view_audit_log", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_audit_log required"
            )
        # Can only see their own client's logs
        query["client_id"] = current_user["client_id"]
    
    # Apply filters
    if user_id:
        query["user_id"] = user_id
    if action_type:
        query["action_type"] = action_type
    if entity_type:
        query["entity_type"] = entity_type
    if from_date:
        query["timestamp"] = {"$gte": from_date}
    if to_date:
        query.setdefault("timestamp", {})["$lte"] = to_date
    
    # Get logs
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("timestamp", -1).skip(skip).limit(limit).to_list(limit)
    
    return [AuditLogEntry(**log) for log in logs]

@api_router.get("/governance/audit/export")
async def export_audit_logs_csv(
    client_id: Optional[str] = None,
    from_date: Optional[str] = None,
    to_date: Optional[str] = None,
    action_type: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Export audit logs as CSV"""
    import csv
    from io import StringIO
    from fastapi.responses import StreamingResponse
    
    # Permission check (same as get_audit_logs)
    query = {}
    if current_user["role"] in ["admin", "recruiter"]:
        if client_id:
            query["client_id"] = client_id
    else:
        has_permission = await check_permission(current_user, "can_export_reports", current_user.get("client_id"))
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_export_reports required"
            )
        query["client_id"] = current_user["client_id"]
    
    # Apply filters
    if action_type:
        query["action_type"] = action_type
    if from_date:
        query["timestamp"] = {"$gte": from_date}
    if to_date:
        query.setdefault("timestamp", {})["$lte"] = to_date
    
    # Get logs
    logs = await db.audit_logs.find(query, {"_id": 0}).sort("timestamp", -1).to_list(10000)
    
    # Create CSV
    output = StringIO()
    writer = csv.DictWriter(output, fieldnames=[
        "log_id", "timestamp", "user_email", "user_role", "client_id",
        "action_type", "entity_type", "entity_id", "previous_value", "new_value"
    ])
    writer.writeheader()
    
    for log in logs:
        writer.writerow({
            "log_id": log.get("log_id"),
            "timestamp": log.get("timestamp"),
            "user_email": log.get("user_email"),
            "user_role": log.get("user_role"),
            "client_id": log.get("client_id"),
            "action_type": log.get("action_type"),
            "entity_type": log.get("entity_type"),
            "entity_id": log.get("entity_id"),
            "previous_value": json.dumps(log.get("previous_value")) if log.get("previous_value") else "",
            "new_value": json.dumps(log.get("new_value")) if log.get("new_value") else ""
        })
    
    output.seek(0)
    
    # Log the export action
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="AUDIT_LOG_EXPORT",
        entity_type="audit_log",
        client_id=client_id,
        metadata={"count": len(logs), "filters": {"from_date": from_date, "to_date": to_date, "action_type": action_type}}
    )
    
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=audit_logs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"}
    )

# Access Matrix Endpoint

@api_router.get("/governance/access-matrix", response_model=list[UserPermissionMatrix])
async def get_access_matrix(
    client_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get access permission matrix for a client"""
    # Permission check
    if current_user["role"] not in ["admin", "recruiter"]:
        # Check if user has permission to view their own client's matrix
        if client_id != current_user.get("client_id"):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
        has_permission = await check_permission(current_user, "can_view_audit_log", client_id)
        if not has_permission:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Permission denied: can_view_audit_log required"
            )
    
    # Get client
    client = await db.clients.find_one({"client_id": client_id}, {"_id": 0})
    if not client:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Client not found"
        )
    
    # Get all users for this client
    users = await db.users.find({"client_id": client_id}, {"_id": 0}).to_list(1000)
    
    matrix = []
    for user in users:
        # Get user's permissions
        perms = await get_user_permissions(user, client_id)
        
        # Get user's roles
        role_assignments = await db.user_client_roles.find({
            "user_id": user.get("user_id", user["email"]),
            "client_id": client_id
        }, {"_id": 0}).to_list(100)
        
        roles = [a["role_name"] for a in role_assignments]
        
        matrix.append(UserPermissionMatrix(
            user_id=user.get("user_id", user["email"]),
            user_email=user["email"],
            user_name=user.get("name", user["email"]),
            client_id=client_id,
            client_name=client["company_name"],
            roles=roles,
            permissions=perms
        ))
    
    return matrix

@api_router.get("/governance/access-matrix/export")
async def export_access_matrix_csv(
    client_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Export access matrix as CSV"""
    import csv
    from io import StringIO
    from fastapi.responses import StreamingResponse
    
    # Get matrix data
    matrix = await get_access_matrix(client_id, current_user)
    
    # Create CSV
    output = StringIO()
    fieldnames = [
        "user_email", "user_name", "client_name", "roles",
        "can_view_jobs", "can_create_jobs", "can_edit_jobs", "can_delete_jobs",
        "can_view_candidates", "can_create_candidates", "can_edit_candidates", "can_delete_candidates",
        "can_update_candidate_status", "can_upload_cv", "can_replace_cv", "can_regenerate_story",
        "can_view_full_cv", "can_view_redacted_cv",
        "can_view_audit_log", "can_manage_roles", "can_manage_users", "can_export_reports"
    ]
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    
    for item in matrix:
        writer.writerow({
            "user_email": item.user_email,
            "user_name": item.user_name,
            "client_name": item.client_name,
            "roles": ", ".join(item.roles),
            "can_view_jobs": item.permissions.can_view_jobs,
            "can_create_jobs": item.permissions.can_create_jobs,
            "can_edit_jobs": item.permissions.can_edit_jobs,
            "can_delete_jobs": item.permissions.can_delete_jobs,
            "can_view_candidates": item.permissions.can_view_candidates,
            "can_create_candidates": item.permissions.can_create_candidates,
            "can_edit_candidates": item.permissions.can_edit_candidates,
            "can_delete_candidates": item.permissions.can_delete_candidates,
            "can_update_candidate_status": item.permissions.can_update_candidate_status,
            "can_upload_cv": item.permissions.can_upload_cv,
            "can_replace_cv": item.permissions.can_replace_cv,
            "can_regenerate_story": item.permissions.can_regenerate_story,
            "can_view_full_cv": item.permissions.can_view_full_cv,
            "can_view_redacted_cv": item.permissions.can_view_redacted_cv,
            "can_view_audit_log": item.permissions.can_view_audit_log,
            "can_manage_roles": item.permissions.can_manage_roles,
            "can_manage_users": item.permissions.can_manage_users,
            "can_export_reports": item.permissions.can_export_reports
        })
    
    output.seek(0)
    
    # Log the export action
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="ACCESS_MATRIX_EXPORT",
        entity_type="access_matrix",
        client_id=client_id,
        metadata={"user_count": len(matrix)}
    )
    
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename=access_matrix_{client_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"}
    )


# ============ INTERVIEW ORCHESTRATION ENDPOINTS ============

@api_router.post("/interviews", response_model=InterviewResponse)
async def create_interview(
    interview_data: InterviewCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a new interview with proposed time slots (Client/Recruiter action)"""
    
    # Verify job exists and get job details
    job = await db.jobs.find_one({"job_id": interview_data.job_id}, {"_id": 0})
    if not job:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found"
        )
    
    # Verify candidate exists
    candidate = await db.candidates.find_one({"candidate_id": interview_data.candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Verify candidate belongs to this job
    if candidate["job_id"] != interview_data.job_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Candidate does not belong to this job"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Get client info
    client = await db.clients.find_one({"client_id": job["client_id"]}, {"_id": 0})
    
    # Generate interview_id
    interview_id = f"int_{uuid.uuid4().hex[:12]}"
    
    # Process proposed slots
    processed_slots = []
    for i, slot in enumerate(interview_data.proposed_slots):
        slot_id = f"slot_{uuid.uuid4().hex[:8]}"
        processed_slots.append({
            "slot_id": slot_id,
            "start_time": slot.get("start_time"),
            "end_time": slot.get("end_time"),
            "duration_minutes": interview_data.interview_duration,
            "is_available": True
        })
    
    now = datetime.now(timezone.utc).isoformat()
    
    interview_doc = {
        "interview_id": interview_id,
        "job_id": interview_data.job_id,
        "candidate_id": interview_data.candidate_id,
        "client_id": job["client_id"],
        "interview_mode": interview_data.interview_mode,
        "interview_duration": interview_data.interview_duration,
        "time_zone": interview_data.time_zone,
        "proposed_slots": processed_slots,
        "selected_slot_id": None,
        "scheduled_start_time": None,
        "scheduled_end_time": None,
        "interview_status": "Awaiting Candidate Confirmation",
        "meeting_link": interview_data.meeting_link,
        "additional_instructions": interview_data.additional_instructions,
        "invite_sent": False,
        "invite_sent_by": None,
        "candidate_confirmation_timestamp": None,
        "no_show_flag": False,
        "no_show_count": 0,
        "interview_round": interview_data.interview_round,
        "round_name": interview_data.round_name or f"Round {interview_data.interview_round}",
        "feedback": None,
        "rating": None,
        "created_at": now,
        "updated_at": now,
        "created_by": current_user["email"]
    }
    
    await db.interviews.insert_one(interview_doc)
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_CREATE",
        entity_type="interview",
        entity_id=interview_id,
        client_id=job["client_id"],
        new_value={
            "candidate_id": interview_data.candidate_id,
            "job_id": interview_data.job_id,
            "interview_mode": interview_data.interview_mode,
            "slots_count": len(processed_slots),
            "round": interview_data.interview_round
        }
    )
    
    return InterviewResponse(
        interview_id=interview_id,
        job_id=interview_data.job_id,
        candidate_id=interview_data.candidate_id,
        client_id=job["client_id"],
        candidate_name=candidate.get("name"),
        job_title=job.get("title"),
        company_name=client.get("company_name") if client else None,
        interview_mode=interview_data.interview_mode,
        interview_duration=interview_data.interview_duration,
        scheduled_start_time=None,
        scheduled_end_time=None,
        time_zone=interview_data.time_zone,
        interview_status="Awaiting Candidate Confirmation",
        meeting_link=interview_data.meeting_link,
        additional_instructions=interview_data.additional_instructions,
        invite_sent=False,
        invite_sent_by=None,
        candidate_confirmation_timestamp=None,
        no_show_flag=False,
        no_show_count=0,
        proposed_slots=processed_slots,
        selected_slot_id=None,
        created_at=now,
        updated_at=now,
        created_by=current_user["email"],
        interview_round=interview_data.interview_round,
        round_name=interview_data.round_name or f"Round {interview_data.interview_round}"
    )


@api_router.get("/interviews", response_model=List[InterviewListItem])
async def list_interviews(
    job_id: Optional[str] = None,
    candidate_id: Optional[str] = None,
    status_filter: Optional[str] = None,
    skip: int = 0,
    limit: int = 100,
    current_user: dict = Depends(get_current_user)
):
    """List interviews with optional filters"""
    query = {}
    
    # Tenant filtering for client users
    if current_user["role"] == "client_user":
        query["client_id"] = current_user["client_id"]
    
    if job_id:
        query["job_id"] = job_id
    if candidate_id:
        query["candidate_id"] = candidate_id
    if status_filter:
        query["interview_status"] = status_filter
    
    interviews = await db.interviews.find(query, {"_id": 0}).skip(skip).limit(limit).sort("created_at", -1).to_list(limit)
    
    result = []
    for interview in interviews:
        # Get candidate and job names
        candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0, "name": 1})
        job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0, "title": 1})
        
        result.append(InterviewListItem(
            interview_id=interview["interview_id"],
            job_id=interview["job_id"],
            candidate_id=interview["candidate_id"],
            candidate_name=candidate.get("name") if candidate else None,
            job_title=job.get("title") if job else None,
            interview_mode=interview["interview_mode"],
            interview_status=interview["interview_status"],
            scheduled_start_time=interview.get("scheduled_start_time"),
            created_at=interview["created_at"]
        ))
    
    return result


@api_router.get("/interviews/{interview_id}", response_model=InterviewResponse)
async def get_interview(
    interview_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get interview details"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Get related info
    candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
    client = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
    
    return InterviewResponse(
        interview_id=interview["interview_id"],
        job_id=interview["job_id"],
        candidate_id=interview["candidate_id"],
        client_id=interview["client_id"],
        candidate_name=candidate.get("name") if candidate else None,
        job_title=job.get("title") if job else None,
        company_name=client.get("company_name") if client else None,
        interview_mode=interview["interview_mode"],
        interview_duration=interview["interview_duration"],
        scheduled_start_time=interview.get("scheduled_start_time"),
        scheduled_end_time=interview.get("scheduled_end_time"),
        time_zone=interview["time_zone"],
        interview_status=interview["interview_status"],
        meeting_link=interview.get("meeting_link"),
        additional_instructions=interview.get("additional_instructions"),
        invite_sent=interview.get("invite_sent", False),
        invite_sent_by=interview.get("invite_sent_by"),
        candidate_confirmation_timestamp=interview.get("candidate_confirmation_timestamp"),
        no_show_flag=interview.get("no_show_flag", False),
        no_show_count=interview.get("no_show_count", 0),
        proposed_slots=interview.get("proposed_slots", []),
        selected_slot_id=interview.get("selected_slot_id"),
        created_at=interview["created_at"],
        updated_at=interview["updated_at"],
        created_by=interview["created_by"]
    )


@api_router.put("/interviews/{interview_id}", response_model=InterviewResponse)
async def update_interview(
    interview_id: str,
    update_data: InterviewUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update interview details"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    update_dict = update_data.model_dump(exclude_unset=True)
    update_dict["updated_at"] = datetime.now(timezone.utc).isoformat()
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_dict}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_UPDATE",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        previous_value={"status": interview["interview_status"]},
        new_value=update_dict
    )
    
    return await get_interview(interview_id, current_user)


@api_router.post("/interviews/{interview_id}/book-slot", response_model=InterviewResponse)
async def book_interview_slot(
    interview_id: str,
    slot_selection: CandidateSlotSelection,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Candidate books an interview slot"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Verify interview is in correct status
    if interview["interview_status"] != "Awaiting Candidate Confirmation":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Interview is not awaiting confirmation (current status: {interview['interview_status']})"
        )
    
    # Find the selected slot
    selected_slot = None
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_selection.slot_id:
            selected_slot = slot
            break
    
    if not selected_slot:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Slot not found"
        )
    
    if not selected_slot.get("is_available", True):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Slot is no longer available"
        )
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update interview with selected slot
    update_data = {
        "selected_slot_id": slot_selection.slot_id,
        "scheduled_start_time": selected_slot["start_time"],
        "scheduled_end_time": selected_slot["end_time"],
        "interview_status": "Confirmed" if slot_selection.confirmed else "Awaiting Candidate Confirmation",
        "candidate_confirmation_timestamp": now if slot_selection.confirmed else None,
        "updated_at": now
    }
    
    # Mark selected slot as unavailable
    updated_slots = []
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_selection.slot_id:
            slot["is_available"] = False
        updated_slots.append(slot)
    update_data["proposed_slots"] = updated_slots
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_data}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_SLOT_BOOKED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        new_value={
            "slot_id": slot_selection.slot_id,
            "start_time": selected_slot["start_time"],
            "confirmed": slot_selection.confirmed
        }
    )
    
    # Send notification if confirmed
    if slot_selection.confirmed:
        slot_time = selected_slot.get("start_time", "TBD")
        background_tasks.add_task(
            send_interview_booking_notification,
            interview_id=interview_id,
            candidate_id=interview["candidate_id"],
            slot_time=slot_time,
            booked_by=current_user["email"]
        )
    
    return await get_interview(interview_id, current_user)


class SendInterviewInviteRequest(BaseModel):
    """Request to send interview invitation email"""
    meeting_link: Optional[str] = None
    interview_mode: Optional[str] = "Video"
    duration_minutes: Optional[int] = 30
    time_zone: Optional[str] = "IST (Indian Standard Time)"
    auto_create_calendar_event: Optional[bool] = False


@api_router.post("/interviews/{interview_id}/send-invite")
async def send_interview_invite(
    interview_id: str,
    request: SendInterviewInviteRequest = None,
    current_user: dict = Depends(get_current_user)
):
    """Send interview invitation email to candidate"""
    from notification_service import send_email, get_interview_invitation_email_template, create_google_calendar_event
    
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    # Get candidate details
    candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    candidate_email = candidate.get("email")
    if not candidate_email:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Candidate does not have an email address"
        )
    
    # Get job and client details
    job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
    client = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
    
    if not job or not client:
        raise HTTPException(status_code=404, detail="Job or client not found")
    
    # Get recruiter info
    recruiter = await db.users.find_one({"email": current_user["email"]}, {"_id": 0, "password_hash": 0})
    
    # Prepare interview data
    interview_data = {
        "interview_id": interview_id,
        "scheduled_at": interview.get("scheduled_start_time", ""),
        "duration_minutes": request.duration_minutes if request else 30,
        "interview_mode": request.interview_mode if request else "Video",
        "time_zone": request.time_zone if request else "IST (Indian Standard Time)",
        "meeting_link": request.meeting_link if request else ""
    }
    
    # Auto-create Google Calendar event if requested
    calendar_result = None
    if request and request.auto_create_calendar_event:
        calendar_result = await create_google_calendar_event(
            interview_data,
            candidate,
            job,
            client
        )
        if calendar_result.get("success"):
            interview_data["meeting_link"] = calendar_result.get("meeting_link", "")
    
    # Generate and send email
    subject, body = get_interview_invitation_email_template(
        candidate=candidate,
        job=job,
        client=client,
        interview=interview_data,
        recruiter=recruiter
    )
    
    email_result = await send_email(candidate_email, subject, body)
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update interview record
    update_data = {
        "invite_sent": True,
        "invite_sent_by": current_user["email"],
        "invite_sent_at": now,
        "interview_status": "Scheduled" if interview["interview_status"] == "Confirmed" else interview["interview_status"],
        "meeting_link": interview_data.get("meeting_link", ""),
        "interview_mode": interview_data.get("interview_mode", "Video"),
        "duration_minutes": interview_data.get("duration_minutes", 30),
        "time_zone": interview_data.get("time_zone", ""),
        "updated_at": now
    }
    
    if calendar_result and calendar_result.get("success"):
        update_data["calendar_event_id"] = calendar_result.get("event_id", "")
        update_data["calendar_link"] = calendar_result.get("calendar_link", "")
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_data}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_INVITE_SENT",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        new_value={
            "candidate_email": candidate_email,
            "email_sent": email_result.get("success", False),
            "calendar_created": calendar_result.get("success", False) if calendar_result else False
        }
    )
    
    return {
        "message": "Interview invitation sent successfully",
        "interview_id": interview_id,
        "email_sent": email_result.get("success", False),
        "candidate_email": candidate_email,
        "meeting_link": interview_data.get("meeting_link", ""),
        "calendar_event_created": calendar_result.get("success", False) if calendar_result else False
    }


@api_router.post("/interviews/{interview_id}/mark-completed")
async def mark_interview_completed(
    interview_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Mark interview as completed"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    now = datetime.now(timezone.utc).isoformat()
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": {
            "interview_status": "Completed",
            "updated_at": now
        }}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_COMPLETED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"]
    )
    
    return {"message": "Interview marked as completed", "interview_id": interview_id}


@api_router.post("/interviews/{interview_id}/mark-no-show")
async def mark_interview_no_show(
    interview_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Mark interview as no-show"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Increment no-show count for candidate
    current_no_show_count = interview.get("no_show_count", 0) + 1
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": {
            "interview_status": "No Show",
            "no_show_flag": True,
            "no_show_count": current_no_show_count,
            "updated_at": now
        }}
    )
    
    # Update candidate's no-show count in candidates collection
    await db.candidates.update_one(
        {"candidate_id": interview["candidate_id"]},
        {"$inc": {"no_show_count": 1}}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_NO_SHOW",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        metadata={"no_show_count": current_no_show_count}
    )
    
    return {"message": "Interview marked as no-show", "interview_id": interview_id, "no_show_count": current_no_show_count}


@api_router.post("/interviews/{interview_id}/cancel")
async def cancel_interview(
    interview_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Cancel an interview"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    now = datetime.now(timezone.utc).isoformat()
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": {
            "interview_status": "Cancelled",
            "updated_at": now
        }}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_CANCELLED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"]
    )
    
    return {"message": "Interview cancelled", "interview_id": interview_id}


# ============ MULTI-ROUND INTERVIEW FLOW ============

class MoveToNextRoundRequest(BaseModel):
    """Request to move candidate to next interview round"""
    feedback: Optional[str] = None
    rating: Optional[int] = Field(default=None, ge=1, le=5, description="Rating 1-5")
    next_round_name: Optional[str] = None  # e.g., "Technical Round 2", "HR Round"

class InitiateHiringRequest(BaseModel):
    """Request to initiate hiring process for a candidate"""
    feedback: Optional[str] = None
    rating: Optional[int] = Field(default=None, ge=1, le=5)
    salary_offered: Optional[str] = None
    joining_date: Optional[str] = None
    offer_notes: Optional[str] = None


@api_router.post("/interviews/{interview_id}/move-to-next-round")
async def move_to_next_round(
    interview_id: str,
    request: MoveToNextRoundRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """
    Mark current interview as Passed and enable scheduling of next round.
    The candidate stays in the pipeline and client can schedule the next round.
    """
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    # Verify access
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(status_code=403, detail="Access denied")
    
    # Interview must be completed first
    if interview["interview_status"] not in ["Completed", "Scheduled"]:
        raise HTTPException(
            status_code=400,
            detail=f"Interview must be Completed or Scheduled before moving to next round. Current status: {interview['interview_status']}"
        )
    
    now = datetime.now(timezone.utc).isoformat()
    current_round = interview.get("interview_round", 1)
    
    # Update current interview as Passed
    update_data = {
        "interview_status": "Passed",
        "feedback": request.feedback,
        "rating": request.rating,
        "updated_at": now,
        "passed_by": current_user["email"],
        "passed_at": now
    }
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_data}
    )
    
    # Update candidate status to indicate they're progressing
    await db.candidates.update_one(
        {"candidate_id": interview["candidate_id"]},
        {"$set": {
            "status": "IN_PROGRESS",
            "current_round": current_round + 1,
            "last_interview_passed": interview_id,
            "updated_at": now
        }}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_PASSED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        metadata={
            "round": current_round,
            "feedback": request.feedback,
            "rating": request.rating,
            "next_round": current_round + 1
        }
    )
    
    # Create notification
    candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
    notification_doc = {
        "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
        "type": "interview_passed",
        "title": f"Interview Passed: {candidate.get('name', 'Candidate')} - Round {current_round}",
        "message": f"Ready for Round {current_round + 1}. {request.next_round_name or ''}",
        "entity_type": "interview",
        "entity_id": interview_id,
        "candidate_id": interview["candidate_id"],
        "created_at": now,
        "read": False,
        "recipients": ["admin", "recruiter", interview["client_id"]]
    }
    await db.notifications.insert_one(notification_doc)
    
    return {
        "message": f"Candidate passed Round {current_round}. Ready for Round {current_round + 1}.",
        "interview_id": interview_id,
        "current_round": current_round,
        "next_round": current_round + 1,
        "candidate_id": interview["candidate_id"],
        "status": "Passed"
    }


@api_router.post("/interviews/{interview_id}/reject")
async def reject_after_interview(
    interview_id: str,
    request: MoveToNextRoundRequest,
    current_user: dict = Depends(get_current_user)
):
    """Mark interview as Failed and reject candidate from this job"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    # Verify access
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(status_code=403, detail="Access denied")
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update interview as Failed
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": {
            "interview_status": "Failed",
            "feedback": request.feedback,
            "rating": request.rating,
            "updated_at": now,
            "rejected_by": current_user["email"],
            "rejected_at": now
        }}
    )
    
    # Update candidate status to REJECTED
    await db.candidates.update_one(
        {"candidate_id": interview["candidate_id"]},
        {"$set": {
            "status": "REJECTED",
            "rejection_reason": request.feedback,
            "rejected_at_round": interview.get("interview_round", 1),
            "updated_at": now
        }}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="INTERVIEW_FAILED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        metadata={
            "round": interview.get("interview_round", 1),
            "feedback": request.feedback
        }
    )
    
    return {
        "message": "Candidate rejected",
        "interview_id": interview_id,
        "candidate_id": interview["candidate_id"],
        "status": "Failed"
    }


@api_router.post("/interviews/{interview_id}/initiate-hiring")
async def initiate_hiring(
    interview_id: str,
    request: InitiateHiringRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """
    Initiate the hiring process for a candidate who has passed all interview rounds.
    This marks the candidate as SELECTED and starts the offer process.
    """
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(status_code=404, detail="Interview not found")
    
    # Verify access
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(status_code=403, detail="Access denied")
    
    now = datetime.now(timezone.utc).isoformat()
    current_round = interview.get("interview_round", 1)
    
    # Update interview status
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": {
            "interview_status": "Passed",
            "feedback": request.feedback,
            "rating": request.rating,
            "hiring_initiated": True,
            "hiring_initiated_at": now,
            "hiring_initiated_by": current_user["email"],
            "updated_at": now
        }}
    )
    
    # Get candidate details
    candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
    
    # Update candidate status to SELECTED
    await db.candidates.update_one(
        {"candidate_id": interview["candidate_id"]},
        {"$set": {
            "status": "SELECTED",
            "selected_at": now,
            "selected_by": current_user["email"],
            "salary_offered": request.salary_offered,
            "proposed_joining_date": request.joining_date,
            "offer_notes": request.offer_notes,
            "total_rounds_cleared": current_round,
            "updated_at": now
        }}
    )
    
    # Log audit event
    await log_audit_event(
        user_id=current_user.get("user_id", current_user["email"]),
        user_email=current_user["email"],
        user_role=current_user["role"],
        action_type="HIRING_INITIATED",
        entity_type="candidate",
        entity_id=interview["candidate_id"],
        client_id=interview["client_id"],
        metadata={
            "interview_id": interview_id,
            "rounds_cleared": current_round,
            "salary_offered": request.salary_offered,
            "joining_date": request.joining_date
        }
    )
    
    # Create notification for recruiters
    job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
    notification_doc = {
        "notification_id": f"notif_{uuid.uuid4().hex[:12]}",
        "type": "hiring_initiated",
        "title": f"Hiring Initiated: {candidate.get('name', 'Candidate')}",
        "message": f"Selected for {job.get('title', 'Position')} after {current_round} round(s)",
        "entity_type": "candidate",
        "entity_id": interview["candidate_id"],
        "created_at": now,
        "read": False,
        "recipients": ["admin", "recruiter"]
    }
    await db.notifications.insert_one(notification_doc)
    
    return {
        "message": f"Hiring initiated for candidate after {current_round} round(s)",
        "interview_id": interview_id,
        "candidate_id": interview["candidate_id"],
        "candidate_name": candidate.get("name"),
        "status": "SELECTED",
        "rounds_cleared": current_round,
        "salary_offered": request.salary_offered,
        "joining_date": request.joining_date
    }


@api_router.get("/candidates/{candidate_id}/interview-history")
async def get_candidate_interview_history(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all interview rounds for a candidate"""
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(status_code=404, detail="Candidate not found")
    
    # Get all interviews for this candidate, sorted by round
    interviews = await db.interviews.find(
        {"candidate_id": candidate_id},
        {"_id": 0}
    ).sort([("interview_round", 1), ("created_at", 1)]).to_list(50)
    
    return {
        "candidate_id": candidate_id,
        "candidate_name": candidate.get("name"),
        "candidate_status": candidate.get("status"),
        "total_rounds": len(interviews),
        "current_round": candidate.get("current_round", 1),
        "interviews": [
            {
                "interview_id": i["interview_id"],
                "round": i.get("interview_round", 1),
                "round_name": i.get("round_name", f"Round {i.get('interview_round', 1)}"),
                "status": i["interview_status"],
                "scheduled_time": i.get("scheduled_start_time"),
                "feedback": i.get("feedback"),
                "rating": i.get("rating"),
                "interview_mode": i["interview_mode"]
            }
            for i in interviews
        ]
    }


@api_router.get("/interviews/stats/pipeline", response_model=InterviewPipelineStats)
async def get_interview_pipeline_stats(
    client_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Get interview pipeline statistics for dashboard"""
    query = {}
    
    # Tenant filtering
    if current_user["role"] == "client_user":
        query["client_id"] = current_user["client_id"]
    elif client_id:
        query["client_id"] = client_id
    
    # Get counts by status
    pipeline = [
        {"$match": query},
        {"$group": {
            "_id": "$interview_status",
            "count": {"$sum": 1}
        }}
    ]
    
    status_counts = await db.interviews.aggregate(pipeline).to_list(100)
    
    # Process results
    stats = {
        "total_interviews": 0,
        "awaiting_confirmation": 0,
        "confirmed": 0,
        "scheduled": 0,
        "completed": 0,
        "no_shows": 0,
        "cancelled": 0
    }
    
    status_mapping = {
        "Awaiting Candidate Confirmation": "awaiting_confirmation",
        "Confirmed": "confirmed",
        "Scheduled": "scheduled",
        "Completed": "completed",
        "No Show": "no_shows",
        "Cancelled": "cancelled"
    }
    
    for item in status_counts:
        status_key = status_mapping.get(item["_id"])
        if status_key:
            stats[status_key] = item["count"]
        stats["total_interviews"] += item["count"]
    
    return InterviewPipelineStats(**stats)


@api_router.get("/candidates/{candidate_id}/interviews", response_model=List[InterviewListItem])
async def get_candidate_interviews(
    candidate_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get all interviews for a candidate"""
    # Verify candidate exists
    candidate = await db.candidates.find_one({"candidate_id": candidate_id}, {"_id": 0})
    if not candidate:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Candidate not found"
        )
    
    # Get job and check tenant
    job = await db.jobs.find_one({"job_id": candidate["job_id"]}, {"_id": 0})
    
    if current_user["role"] == "client_user":
        if job and job["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    interviews = await db.interviews.find(
        {"candidate_id": candidate_id},
        {"_id": 0}
    ).sort("created_at", -1).to_list(100)
    
    result = []
    for interview in interviews:
        result.append(InterviewListItem(
            interview_id=interview["interview_id"],
            job_id=interview["job_id"],
            candidate_id=interview["candidate_id"],
            candidate_name=candidate.get("name"),
            job_title=job.get("title") if job else None,
            interview_mode=interview["interview_mode"],
            interview_status=interview["interview_status"],
            scheduled_start_time=interview.get("scheduled_start_time"),
            created_at=interview["created_at"]
        ))
    
    return result


# ============ PUBLIC CANDIDATE BOOKING ENDPOINTS (No Auth Required) ============

def generate_booking_token(interview_id: str) -> str:
    """Generate a simple booking token for candidate access"""
    import hashlib
    secret = os.environ.get('JWT_SECRET', 'arbeit-secret-key')
    return hashlib.sha256(f"{interview_id}:{secret}".encode()).hexdigest()[:32]

def verify_booking_token(interview_id: str, token: str) -> bool:
    """Verify the booking token"""
    expected = generate_booking_token(interview_id)
    return token == expected


@api_router.get("/public/interviews/{interview_id}")
async def get_public_interview(interview_id: str, token: str):
    """Public endpoint for candidates to view interview details"""
    
    if not verify_booking_token(interview_id, token):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Invalid or expired booking link"
        )
    
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Get related info
    candidate = await db.candidates.find_one({"candidate_id": interview["candidate_id"]}, {"_id": 0})
    job = await db.jobs.find_one({"job_id": interview["job_id"]}, {"_id": 0})
    client = await db.clients.find_one({"client_id": interview["client_id"]}, {"_id": 0})
    
    return {
        "interview_id": interview["interview_id"],
        "interview_mode": interview["interview_mode"],
        "interview_duration": interview["interview_duration"],
        "time_zone": interview["time_zone"],
        "interview_status": interview["interview_status"],
        "proposed_slots": interview.get("proposed_slots", []),
        "scheduled_start_time": interview.get("scheduled_start_time"),
        "scheduled_end_time": interview.get("scheduled_end_time"),
        "meeting_link": interview.get("meeting_link"),
        "additional_instructions": interview.get("additional_instructions"),
        "candidate_name": candidate.get("name") if candidate else None,
        "job_title": job.get("title") if job else None,
        "company_name": client.get("company_name") if client else None
    }


@api_router.post("/public/interviews/{interview_id}/book")
async def public_book_slot(interview_id: str, slot_id: str, token: str):
    """Public endpoint for candidates to book a slot"""
    
    if not verify_booking_token(interview_id, token):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Invalid or expired booking link"
        )
    
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Verify interview is in correct status
    if interview["interview_status"] != "Awaiting Candidate Confirmation":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Interview is not awaiting confirmation (current status: {interview['interview_status']})"
        )
    
    # Find the selected slot
    selected_slot = None
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_id:
            selected_slot = slot
            break
    
    if not selected_slot:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Slot not found"
        )
    
    if not selected_slot.get("is_available", True):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Slot is no longer available"
        )
    
    now = datetime.now(timezone.utc).isoformat()
    
    # Update interview with selected slot
    update_data = {
        "selected_slot_id": slot_id,
        "scheduled_start_time": selected_slot["start_time"],
        "scheduled_end_time": selected_slot["end_time"],
        "interview_status": "Confirmed",
        "candidate_confirmation_timestamp": now,
        "updated_at": now
    }
    
    # Mark selected slot as unavailable
    updated_slots = []
    for slot in interview.get("proposed_slots", []):
        if slot["slot_id"] == slot_id:
            slot["is_available"] = False
        updated_slots.append(slot)
    update_data["proposed_slots"] = updated_slots
    
    await db.interviews.update_one(
        {"interview_id": interview_id},
        {"$set": update_data}
    )
    
    # Log audit event (without user context since this is public)
    await log_audit_event(
        user_id="candidate",
        user_email="candidate-booking",
        user_role="candidate",
        action_type="INTERVIEW_SLOT_BOOKED",
        entity_type="interview",
        entity_id=interview_id,
        client_id=interview["client_id"],
        new_value={
            "slot_id": slot_id,
            "start_time": selected_slot["start_time"],
            "booked_via": "public_link"
        }
    )
    
    return {"message": "Interview slot confirmed", "interview_id": interview_id}


@api_router.get("/interviews/{interview_id}/booking-link")
async def get_booking_link(
    interview_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get the booking link for an interview (to share with candidate)"""
    interview = await db.interviews.find_one({"interview_id": interview_id}, {"_id": 0})
    if not interview:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Interview not found"
        )
    
    # Tenant check for client users
    if current_user["role"] == "client_user":
        if interview["client_id"] != current_user["client_id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Access denied"
            )
    
    booking_token = generate_booking_token(interview_id)
    
    # Get the frontend URL from environment or construct it
    frontend_url = os.environ.get('FRONTEND_URL', 'https://hirematch-52.preview.emergentagent.com')
    booking_link = f"{frontend_url}/book/{interview_id}/{booking_token}"
    
    return {
        "interview_id": interview_id,
        "booking_link": booking_link,
        "booking_token": booking_token
    }


# ============ NOTIFICATION ENDPOINTS ============

class NotificationResponse(BaseModel):
    notification_id: str
    type: str
    title: str
    message: str
    entity_type: Optional[str] = None
    entity_id: Optional[str] = None
    client_id: Optional[str] = None
    created_at: str
    created_by: str
    is_read: bool = False


@api_router.get("/notifications", response_model=List[NotificationResponse])
async def get_notifications(
    unread_only: bool = False,
    limit: int = 50,
    current_user: dict = Depends(get_current_user)
):
    """Get notifications for the current user"""
    query = {
        "$or": [
            {"for_roles": current_user["role"]},
            {"for_users": current_user["email"]}
        ]
    }
    
    # For client users, also filter by client_id
    if current_user["role"] == "client_user":
        query["$or"].append({"client_id": current_user.get("client_id")})
    
    notifications = await db.notifications.find(
        query,
        {"_id": 0}
    ).sort("created_at", -1).limit(limit).to_list(limit)
    
    result = []
    for notif in notifications:
        is_read = current_user["email"] in notif.get("read_by", [])
        if unread_only and is_read:
            continue
        result.append(NotificationResponse(
            notification_id=notif["notification_id"],
            type=notif["type"],
            title=notif["title"],
            message=notif["message"],
            entity_type=notif.get("entity_type"),
            entity_id=notif.get("entity_id"),
            client_id=notif.get("client_id"),
            created_at=notif["created_at"],
            created_by=notif["created_by"],
            is_read=is_read
        ))
    
    return result


@api_router.get("/notifications/unread-count")
async def get_unread_notification_count(
    current_user: dict = Depends(get_current_user)
):
    """Get count of unread notifications"""
    query = {
        "$or": [
            {"for_roles": current_user["role"]},
            {"for_users": current_user["email"]}
        ],
        "read_by": {"$ne": current_user["email"]}
    }
    
    if current_user["role"] == "client_user":
        query["$or"].append({"client_id": current_user.get("client_id")})
    
    count = await db.notifications.count_documents(query)
    return {"unread_count": count}


@api_router.post("/notifications/{notification_id}/mark-read")
async def mark_notification_read(
    notification_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Mark a notification as read"""
    result = await db.notifications.update_one(
        {"notification_id": notification_id},
        {"$addToSet": {"read_by": current_user["email"]}}
    )
    
    if result.modified_count == 0:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Notification not found"
        )
    
    return {"message": "Notification marked as read"}


@api_router.post("/notifications/mark-all-read")
async def mark_all_notifications_read(
    current_user: dict = Depends(get_current_user)
):
    """Mark all notifications as read for the current user"""
    query = {
        "$or": [
            {"for_roles": current_user["role"]},
            {"for_users": current_user["email"]}
        ]
    }
    
    result = await db.notifications.update_many(
        query,
        {"$addToSet": {"read_by": current_user["email"]}}
    )
    
    return {"message": f"Marked {result.modified_count} notifications as read"}


@api_router.get("/health")
async def health_check():
    return {"status": "healthy"}

# Include the router in the main app
app.include_router(api_router)

# Mount static files directory for serving uploaded CVs
app.mount("/api/uploads", StaticFiles(directory=str(UPLOAD_DIR)), name="uploads")

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()